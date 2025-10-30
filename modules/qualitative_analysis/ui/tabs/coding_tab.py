"""
Tab: Asistente de Codificación Cualitativa
Sugerencias de códigos con IA, revisión humana y exportación CSV.
"""

import streamlit as st
from typing import List, Dict, Any

from ...core.config import AnalysisConfig
from ...extractors.coding_assistant import QualitativeCodingAssistant, CodeSuggestion
from ..components.educational import (
    show_methodology_box,
    show_interpretation_guide,
    show_statistics_panel,
    show_help_sidebar
)


def render_coding_tab(chunks: List[Dict[str, Any]], config: AnalysisConfig):
    # Encabezado
    st.markdown(
        """
        <div style="
            background: linear-gradient(90deg, #ffcc00 0%, #ff9900 100%);
            padding: 1.5rem;
            border-radius: 10px;
            margin-bottom: 2rem;
            text-align: center;
        ">
            <h1 style="color: #1b1e23; margin: 0; font-size: 2rem;">🛠️ Asistente de Codificación Cualitativa</h1>
            <p style="color: #1b1e23; margin: 0.5rem 0 0 0; font-size: 1.05rem; opacity: 0.9;">
                Sugerencias de códigos con IA + revisión humana y citación
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # Ayuda lateral
    show_help_sidebar()

    # Metodología
    if config.show_methodology:
        show_methodology_box(
            title="Candidatos con n-gramas/TF-IDF + Refinamiento LLM",
            description=(
                "Se generan términos candidatos (palabras/frases frecuentes y distintivas) y se refinan con IA "
                "para proponer nombres de códigos claros, definiciones operacionales y ejemplos coherentes. Cada sugerencia "
                "incluye ejemplos generados por IA y citas a fragmentos del corpus para trazabilidad."
            ),
            steps=[
                "Preprocesamiento y vectorización (TF-IDF con n-gramas)",
                "Selección de candidatos y agrupación básica",
                "Refinamiento con LLM (nombres, definiciones y ejemplos coherentes)",
                "Citación de fragmentos relevantes como evidencia adicional",
                "Revisión humana: aprobar, editar, descartar; exportación a CSV",
            ],
            algorithm_info="TF-IDF (1-3) + DeepSeek R1 (refinamiento de código y generación de ejemplos)"
        )

    if config.show_interpretation_guide:
        show_interpretation_guide(
            what_it_means=(
                "Los códigos representan categorías analíticas operacionales. Las sugerencias sirven como punto de partida; "
                "el investigador valida y ajusta. Cada código incluye ejemplos coherentes generados por IA y citas que "
                "enlazan cada código con su evidencia en los documentos originales."
            ),
            how_to_use=[
                "Revisa la lista, edita nombres/definiciones/ejemplos, y aprueba los códigos útiles",
                "Usa los ejemplos y las citas para validar el significado y coherencia",
                "Los ejemplos del LLM representan el concepto de forma sintética",
                "Exporta el libro de códigos inicial a CSV (NVivo/Atlas.ti)"
            ],
            limitations=[
                "Las sugerencias requieren validación humana",
                "El LLM genera ejemplos sintéticos que deben validarse con las citas",
                "Los nombres pueden ser útiles aunque no aparezcan literalmente en el texto",
                "Las definiciones y ejemplos deben adaptarse al proyecto específico"
            ],
        )

    if not chunks:
        st.warning("⚠️ No hay documentos procesados para este análisis.")
        return

    st.divider()

    # Estadísticas de entrada (coherente con otros sub-módulos)
    st.markdown("### 🧾 Estadísticas de Entrada")
    sources = list(set(ch.get('metadata', {}).get('source_file', 'unknown') for ch in chunks))
    total_content = sum(len(ch.get('content', '')) for ch in chunks)

    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("📄 Documentos", len(sources))
    with col2:
        st.metric("📝 Chunks", len(chunks))
    with col3:
        st.metric("📊 Caracteres", f"{total_content:,}")

    with st.expander("📋 Ver lista de fuentes"):
        for i, src in enumerate(sources, 1):
            st.markdown(f"{i}. `{src}`")

    # Configuración del análisis (valores predefinidos en el código)
    st.markdown("### ⚙️ Configuración del Análisis")
    st.info(
        """
        **ℹ️ Configuración Automática:** Los parámetros del asistente de codificación están optimizados automáticamente.
        Si necesitas modificar algún parámetro específico, puedes hacerlo directamente en el código (`AnalysisConfig`).
        """
    )

    st.divider()

    # Botones (patrón consistente con otros sub-módulos)
    col1, col2, col3 = st.columns([2, 1, 2])
    with col1:
        if st.session_state.get('coding_analyzed', False):
            if st.button("🔄 Nuevo Análisis", type="secondary", use_container_width=True, key="coding_new_analysis", help="Limpiar resultados y realizar nuevo análisis"):
                # Limpiar estado del asistente de codificación
                st.session_state.coding_analyzed = False
                st.session_state.coding_suggestions = []
                st.session_state.coding_approved = []
                if 'coding_drafts' in st.session_state:
                    del st.session_state['coding_drafts']
                if 'coding_assistant' in st.session_state:
                    del st.session_state['coding_assistant']
                st.rerun()
    with col2:
        if not st.session_state.get('coding_analyzed', False):
            analyze_button = st.button(
                "🚀 Sugerir Códigos",
                type="primary",
                use_container_width=True,
                key="coding_suggest"
            )
        else:
            analyze_button = False
    with col3:
        if st.session_state.get('coding_analyzed', False):
            st.success("✅ Análisis completado")

    # Session state
    if 'coding_suggestions' not in st.session_state:
        st.session_state.coding_suggestions: List[CodeSuggestion] = []
    if 'coding_approved' not in st.session_state:
        st.session_state.coding_approved: List[Dict[str, Any]] = []

    # Ejecutar
    if analyze_button:
        with st.spinner("🔍 Generando sugerencias de códigos..."):
            try:
                assistant = QualitativeCodingAssistant(config)
                suggestions = assistant.suggest_codes(chunks)

                st.session_state.coding_suggestions = suggestions
                # Limpiar aprobados previos al nuevo análisis
                st.session_state.coding_approved = []
                # Guardar assistant para futura exportación de citas si se quisiera
                st.session_state.coding_assistant = assistant
                st.session_state.coding_analyzed = True
                st.success(f"✅ {len(suggestions)} códigos sugeridos")
                st.rerun()
            except Exception as e:
                st.error(f"❌ Error generando sugerencias: {str(e)}")
                st.exception(e)
                return

    # Mostrar sugerencias y exportación en pestañas sólo cuando hay análisis
    suggestions: List[CodeSuggestion] = st.session_state.get('coding_suggestions', [])
    if st.session_state.get('coding_analyzed', False) and suggestions:
        # Inicializar borradores editables en session_state para evitar perder cambios con los reruns
        if 'coding_drafts' not in st.session_state or not isinstance(st.session_state.coding_drafts, list) or len(st.session_state.coding_drafts) != len(suggestions):
            drafts: List[Dict[str, Any]] = []
            for s in suggestions:
                # Preparar citas iniciales en formato editable
                init_citations = []
                if getattr(s, 'citations', None):
                    for c in s.citations[:3]:
                        init_citations.append({
                            'source': getattr(c, 'source_file', 'unknown'),
                            'text': c.get_full_context(max_chars=220)
                        })
                drafts.append({
                    'name': s.code,
                    'definition': s.definition,
                    'example': s.examples[0] if s.examples else "",
                    'citations': init_citations
                })
            st.session_state.coding_drafts = drafts

        tab1, tab2 = st.tabs(["🎯 Sugerencias", "💾 Exportar"])

        with tab1:
            st.markdown("### 🎯 Sugerencias de Códigos")
            col_a, col_b = st.columns([3, 1])
            with col_a:
                filt = st.text_input("Filtrar por texto", key="coding_filter_text")
            with col_b:
                order = st.selectbox("Ordenar por", ["Confianza", "Nombre"], key="coding_sort_by")

            filtered = suggestions
            if filt:
                filtered = [s for s in filtered if filt.lower() in s.code.lower() or filt.lower() in s.definition.lower()]
            if order == "Nombre":
                filtered = sorted(filtered, key=lambda s: s.code.lower())
            else:
                filtered = sorted(filtered, key=lambda s: s.confidence, reverse=True)

            # Usamos índices de las sugerencias completas para sincronizar con drafts
            for idx, s in enumerate(suggestions):
                i = idx + 1
                draft = st.session_state.coding_drafts[idx]
                with st.expander(f"{i}. {draft['name']}", expanded=False):
                    # Campos editables vinculados al borrador persistente
                    draft['name'] = st.text_input("Nombre del código", value=draft['name'], key=f"coding_name_{idx}")
                    draft['definition'] = st.text_area("Definición", value=draft['definition'], key=f"coding_def_{idx}")

                    draft['example'] = st.text_area(
                        "Ejemplo (editable)",
                        value=draft['example'],
                        height=120,
                        key=f"coding_example_{idx}"
                    )

                    # Citas editables con agregar/quitar
                    st.markdown("**Citas (editables, 0 a 5):**")
                    if 'citations' not in draft or not isinstance(draft['citations'], list):
                        draft['citations'] = []
                    to_remove = []
                    for j, cit in enumerate(draft['citations'][:5]):
                        with st.expander(f"Cita {j+1}", expanded=False):
                            cit['source'] = st.text_input(
                                f"Fuente (documento) {j+1}",
                                value=cit.get('source', 'unknown'),
                                key=f"coding_citation_source_{idx}_{j}"
                            )
                            cit['text'] = st.text_area(
                                f"Texto de la cita {j+1}",
                                value=cit.get('text', ''),
                                height=100,
                                key=f"coding_citation_text_{idx}_{j}"
                            )
                            if st.button("🗑️ Quitar cita", key=f"coding_citation_remove_{idx}_{j}"):
                                to_remove.append(j)
                    # Eliminar seleccionadas
                    for j in reversed(to_remove):
                        del draft['citations'][j]
                    # Agregar nueva cita
                    if len(draft['citations']) < 5:
                        if st.button("➕ Agregar cita", key=f"coding_citation_add_{idx}"):
                            draft['citations'].append({'source': 'unknown', 'text': ''})

                    cols = st.columns([1,1,2])
                    with cols[0]:
                        if st.button("✅ Aprobar", key=f"approve_{i}"):
                            st.session_state.coding_approved.append({
                                'code': draft['name'],
                                'definition': draft['definition'],
                                'examples': [draft['example']] if draft['example'] else [],
                                'citations': draft['citations']
                            })
                            st.success("Aprobado y agregado al libro de códigos")
                    with cols[1]:
                        if st.button("🗑️ Descartar", key=f"discard_{i}"):
                            # Eliminar tanto la sugerencia como su borrador asociado
                            st.session_state.coding_suggestions.pop(idx)
                            st.session_state.coding_drafts.pop(idx)
                            st.info("Descartado")

            st.divider()
            st.markdown("### 📘 Libro de Códigos (Aprobados)")
            approved = st.session_state.get('coding_approved', [])
            if not approved:
                st.info("Aún no has aprobado códigos. Usa los botones en cada sugerencia.")
            else:
                st.success(f"{len(approved)} código(s) en el libro inicial")
                st.dataframe(approved, use_container_width=True)

        with tab2:
            st.markdown("### 💾 Exportar Libro de Códigos")
            # Información del documento
            st.markdown("#### Información del Documento")
            colx, coly = st.columns(2)
            with colx:
                doc_title = st.text_input("Título del documento", value="Libro de Códigos", key="coding_export_title")
            with coly:
                doc_author = st.text_input("Autor", value="Investigador", key="coding_export_author")
            approved = st.session_state.get('coding_approved', [])
            if not approved:
                st.info("Aprueba al menos un código para habilitar la exportación.")
            else:
                # CSV
                import csv, io
                buffer = io.StringIO()
                writer = csv.writer(buffer)
                writer.writerow(["Code", "Definition", "Example", "Citations"])
                for item in approved:
                    example = (item.get('examples') or [''])[0] if item.get('examples') else ''
                    cits = item.get('citations') or []
                    cits_join = " | ".join([f"{c.get('source','unknown')}: {c.get('text','')}" for c in cits])
                    writer.writerow([item.get('code',''), item.get('definition',''), example, cits_join])
                csv_bytes = buffer.getvalue().encode('utf-8')

                st.download_button(
                    label="📄 Descargar CSV (NVivo/Atlas.ti)",
                    data=csv_bytes,
                    file_name="codebook.csv",
                    mime="text/csv"
                )

                # Word
                if st.button("📄 Generar Documento Word", type="primary", key="coding_generate_word"):
                    try:
                        doc_bytes = _generate_codebook_word(approved, title=doc_title, author=doc_author)
                        st.download_button(
                            label="💾 Descargar Documento Word",
                            data=doc_bytes,
                            file_name="codebook.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        st.success("✅ Documento Word generado correctamente")
                    except Exception as e:
                        st.error(f"❌ Error al generar documento Word: {str(e)}")


def _generate_codebook_word(approved_codes: List[Dict[str, Any]], title: str = "Libro de Códigos", author: str = "Investigador") -> bytes:
    """
    Generar documento Word con el libro de códigos aprobado.
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io
    except ImportError:
        raise ImportError("python-docx es requerido para exportar a Word. Instala con: pip install python-docx")

    doc = Document()

    # Título
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Autor y fecha
    from datetime import datetime
    doc.add_paragraph(f"Autor: {author}")
    doc.add_paragraph(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph("")

    for i, item in enumerate(approved_codes, 1):
        name = item.get('code', f'Código {i}')
        definition = item.get('definition', '')
        example = (item.get('examples') or [''])[0] if item.get('examples') else ''

        doc.add_heading(f"{i}. {name}", level=1)
        doc.add_paragraph(f"Definición: {definition}")
        if example:
            doc.add_paragraph("Ejemplo:")
            doc.add_paragraph(example)
        # Citas
        citations = item.get('citations') or []
        if citations:
            doc.add_paragraph("Referencias:")
            for k, cit in enumerate(citations, 1):
                doc.add_paragraph(f"{k}. {cit.get('source','unknown')} — {cit.get('text','')}")
        doc.add_paragraph("")

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()