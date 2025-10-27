"""
Tab de Análisis de Sentimientos
Interfaz de usuario para análisis de polaridad y subjetividad
"""

import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from typing import List, Dict, Any, Optional
from datetime import datetime

from ...extractors.sentiment_extractor import SentimentExtractor, SentimentAnalysisResult
from ...core.config import AnalysisConfig
from ..components.educational import (
    show_methodology_box,
    show_interpretation_guide,
    show_citation_box,
    show_statistics_panel
)


def render_sentiment_tab(chunks: List[Dict[str, Any]], config: AnalysisConfig):
    """
    Renderizar tab de análisis de sentimientos
    
    Args:
        chunks: Lista de chunks de texto para análisis
        config: Configuración del análisis
    """
    
    # Título con descripción
    st.markdown("""
    <div style="
        background: linear-gradient(90deg, #e74c3c 0%, #c0392b 100%);
        padding: 1.5rem;
        border-radius: 10px;
        margin-bottom: 2rem;
        text-align: center;
    ">
        <h1 style="color: white; margin: 0; font-size: 2rem;">😊 Análisis de Sentimientos</h1>
        <p style="color: white; margin: 0.5rem 0 0 0; font-size: 1.1rem; opacity: 0.9;">
            Identifica emociones profundas y patrones sentimentales con fundamentación completa
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # Mostrar metodología
    show_methodology_box(
        title="Enfoque Híbrido: Algoritmo + DeepSeek R1",
        description="""
        **¿Qué hace este análisis?**
        
        Este módulo identifica automáticamente los sentimientos y emociones presentes en tus documentos 
        utilizando un **enfoque híbrido** que combina algoritmos avanzados de análisis de sentimientos 
        con inteligencia artificial para generar interpretaciones emocionales profundas y análisis contextuales.
        """,
        steps=[
            "Preprocesamiento: Limpieza y normalización del texto",
            "Vectorización: Conversión a representación numérica (TF-IDF)",
            "Algoritmo de Análisis: VADER, TextBlob o Híbrido para extracción inicial",
            "🤖 Refinamiento IA: DeepSeek R1 analiza y mejora las interpretaciones sentimentales",
            "Generación de análisis emocionales profundos: Crea interpretaciones académicas significativas",
            "Validación: Cálculo de polaridad mejorada y métricas de confianza",
            "Citación: Vinculación con fuentes originales",
            "Visualización: Análisis temporal y distribución emocional"
        ],
        algorithm_info="VADER/TextBlob/Híbrido + DeepSeek R1 para análisis emocionales profundos"
    )
    
    # Mostrar guía de interpretación
    show_interpretation_guide(
        what_it_means="""
        **¿Qué significan estos resultados?**
        
        Los análisis de sentimientos identificados son **interpretaciones emocionales profundas** que emergen de tus documentos. 
        Con el refinamiento IA, cada análisis captura emociones específicas, intensidades y patrones emocionales 
        complejos que van más allá de simples clasificaciones positivo/negativo, proporcionando comprensión matizada 
        del contenido emocional presente en tu investigación.
        """,
        how_to_use=[
            "Identifica los patrones emocionales centrales de tu investigación",
            "Analiza las interpretaciones generadas por la IA para cada sentimiento",
            "Examina las emociones específicas para entender matices emocionales",
            "Estudia los patrones emocionales para entender estructuras sentimentales",
            "Usa las métricas de intensidad y confianza para validar la calidad",
            "Aprovecha los análisis emocionales profundos para desarrollar marcos teóricos",
            "Exporta los resultados para fundamentar tu análisis cualitativo"
        ],
        limitations=[
            "Los análisis generados por IA requieren validación del investigador",
            "La calidad depende de la cantidad y diversidad del contenido analizado",
            "El LLM puede generar interpretaciones que sintetizan información no explícita",
            "Las interpretaciones son análisis que deben ser verificadas",
            "Requiere acceso a un modelo LLM (Ollama) para funcionar completamente"
        ]
    )
    
    # Verificar si hay chunks disponibles
    if not chunks:
        st.warning("⚠️ No hay documentos disponibles para análisis de sentimientos")
        st.info("""
        **Para realizar análisis de sentimientos:**
        1. Asegúrate de que hay documentos procesados en el cache RAG
        2. Los documentos deben contener texto suficiente para análisis
        3. Mínimo recomendado: 5 documentos o 500 palabras
        """)
        return
    
    # Mostrar estadísticas de entrada
    st.markdown("#### 📊 Estadísticas de Entrada")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric(
            "📄 Documentos",
            len(set(chunk.get('metadata', {}).get('source_file', 'unknown') 
                   for chunk in chunks)),
            help="Número de documentos únicos"
        )
    
    with col2:
        st.metric(
            "📝 Chunks",
            len(chunks),
            help="Número total de fragmentos de texto"
        )
    
    with col3:
        st.metric(
            "📖 Palabras Estimadas",
            sum(len(chunk.get('content', '').split()) for chunk in chunks),
            help="Número aproximado de palabras"
        )
    
    with col4:
        st.metric(
            "📂 Fuentes Únicas",
            len(set(chunk.get('metadata', {}).get('source_file', 'unknown') 
                   for chunk in chunks)),
            help="Número de fuentes diferentes"
        )
    
    # Configuración del análisis
    st.markdown("#### ⚙️ Configuración del Análisis")
    
    st.info("""
    **ℹ️ Configuración Automática:** Los parámetros del análisis de sentimientos están optimizados automáticamente.
    Si necesitas modificar algún parámetro específico, puedes hacerlo directamente en el código.
    """)
    
    st.divider()
    
    # Botones de análisis
    col1, col2, col3 = st.columns([2, 1, 2])
    
    with col1:
        if st.session_state.get('sentiments_analyzed', False):
            if st.button("🔄 Nuevo Análisis", type="secondary", use_container_width=True, help="Limpiar resultados y realizar nuevo análisis", key="sentiment_new_analysis"):
                # Limpiar session state
                st.session_state.sentiments_analyzed = False
                st.session_state.sentiment_analysis_result = None
                st.session_state.sentiment_extractor = None
                st.session_state.sentiments_summary = None
                st.rerun()
    
    with col2:
        if not st.session_state.get('sentiments_analyzed', False):
            analyze_button = st.button(
                "🚀 Analizar Sentimientos",
                type="primary",
                use_container_width=True,
                help="Iniciar análisis híbrido de sentimientos",
                key="sentiment_analyze"
            )
        else:
            analyze_button = False
    
    with col3:
        if st.session_state.get('sentiments_analyzed', False):
            st.success("✅ Análisis completado")
    
    # Realizar análisis solo si se presiona el botón Y no hay resultados previos
    if analyze_button and not st.session_state.get('sentiments_analyzed', False):
        with st.spinner("😊 Analizando sentimientos con enfoque híbrido..."):
            try:
                # Crear extractor
                extractor = SentimentExtractor(config)
                
                # Ejecutar análisis híbrido
                result = extractor.analyze_sentiments_hybrid(chunks)
                
                # Guardar en session state
                st.session_state.sentiments_analyzed = True
                st.session_state.sentiment_analysis_result = result
                st.session_state.sentiment_extractor = extractor
                
                # Obtener resumen
                summary = extractor.get_sentiment_summary(result)
                st.session_state.sentiments_summary = summary
                
                # Forzar re-ejecución para mostrar resultados
                st.rerun()
                
            except Exception as e:
                st.error(f"❌ **Error en el análisis:** {str(e)}")
                st.exception(e)
                return
    
    # Mostrar resultados si existen
    if st.session_state.get('sentiments_analyzed', False):
        result = st.session_state['sentiment_analysis_result']
        summary = st.session_state.get('sentiments_summary', {})
        
        render_sentiment_results(result, summary, chunks)


def render_sentiment_results(result: SentimentAnalysisResult, summary: Dict[str, Any], chunks: List[Dict[str, Any]]):
    """
    Renderizar resultados del análisis de sentimientos
    
    Args:
        result: Resultado del análisis de sentimientos
        summary: Resumen estadístico del análisis
        chunks: Chunks originales para referencias
    """
    st.markdown("#### 📊 Resultados del Análisis de Sentimientos")
    
    # Mostrar resumen del análisis
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Textos Analizados", summary.get('total_sentiments', 0))
    
    with col2:
        st.metric("Polaridad Promedio", f"{summary.get('avg_polarity', 0.0):.3f}")
    
    with col3:
        st.metric("Subjetividad Promedio", f"{summary.get('avg_subjectivity', 0.0):.3f}")
    
    with col4:
        st.metric("Algoritmo", summary.get('algorithm_used', 'unknown'))
    
    # Tabs para diferentes vistas de resultados
    tab1, tab2, tab3, tab4 = st.tabs([
        "📋 Sentimientos Principales",
        "📊 Visualizaciones", 
        "🔗 Análisis por Fuente",
        "💾 Exportar"
    ])
    
    with tab1:
        render_sentiments_list(result)
    
    with tab2:
        render_sentiment_visualizations(result)
    
    with tab3:
        render_sentiment_by_source(result)
    
    with tab4:
        render_sentiment_export(result)


def render_sentiments_list(result: SentimentAnalysisResult):
    """Renderizar lista de sentimientos identificados"""
    st.markdown("#### 📋 Sentimientos Principales")
    
    if not result.sentiments:
        st.info("No se encontraron sentimientos para mostrar")
        return
    
    # Crear DataFrame para mostrar sentimientos
    sentiments_data = []
    for i, sentiment in enumerate(result.sentiments):
        sentiments_data.append({
            'ID': i + 1,
            'Texto': sentiment.text[:100] + "..." if len(sentiment.text) > 100 else sentiment.text,
            'Sentimiento': sentiment.sentiment_label,
            'Polaridad': f"{sentiment.polarity:.3f}",
            'Subjetividad': f"{sentiment.subjectivity:.3f}",
            'Confianza': f"{sentiment.confidence:.3f}",
            'Contexto': sentiment.context[:50] + "..." if len(sentiment.context) > 50 else sentiment.context
        })
    
    df = pd.DataFrame(sentiments_data)
    
    # Mostrar tabla con sentimientos
    st.dataframe(df, use_container_width=True)
    
    # Mostrar detalles de cada sentimiento
    st.markdown("#### 🔍 Detalles de Sentimientos")
    
    for i, sentiment in enumerate(result.sentiments[:10]):  # Mostrar solo los primeros 10
        with st.expander(f"**Sentimiento {i+1}: {sentiment.sentiment_label.upper()}**"):
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("**Texto Completo:**")
                st.text_area("Texto", sentiment.text, height=100, disabled=True, label_visibility="collapsed", key=f"sentiment_text_{i}")
                
                st.markdown("**Contexto:**")
                st.info(sentiment.context)
            
            with col2:
                st.markdown("**Métricas:**")
                st.metric("Polaridad", f"{sentiment.polarity:.3f}")
                st.metric("Subjetividad", f"{sentiment.subjectivity:.3f}")
                st.metric("Confianza", f"{sentiment.confidence:.3f}")
                
                # Barra de progreso para visualizar polaridad
                if sentiment.polarity > 0:
                    st.progress(sentiment.polarity, text=f"Positivo: {sentiment.polarity:.1%}")
                else:
                    st.progress(abs(sentiment.polarity), text=f"Negativo: {abs(sentiment.polarity):.1%}")
            
            # Mostrar citaciones si existen
            if sentiment.citations:
                with st.expander("📚 Ver Citaciones"):
                    for j, citation in enumerate(sentiment.citations[:3], 1):
                        st.markdown(f"**Cita {j}:**")
                        st.markdown(f"📄 **Fuente:** {citation.get('source_file', 'unknown')}")
                        st.markdown(f"📝 **Contenido:** {citation.get('content', '')[:200]}...")
                        st.markdown(f"🔗 **Chunk:** {citation.get('chunk_id', 0)}")
                        st.markdown("---")


def render_sentiment_visualizations(result: SentimentAnalysisResult):
    """Renderizar visualizaciones de sentimientos"""
    st.markdown("#### 📊 Visualizaciones de Sentimientos")
    
    # Crear extractor para generar visualizaciones
    extractor = SentimentExtractor(AnalysisConfig())
    
    try:
        # Generar visualizaciones
        viz_data = extractor.generate_sentiment_visualization(result)
        
        if viz_data:
            # Gráfico de distribución de sentimientos
            st.plotly_chart(viz_data['distribution_chart'], use_container_width=True)
            
            # Gráfico de dispersión polaridad vs subjetividad
            st.plotly_chart(viz_data['scatter_chart'], use_container_width=True)
            
            # Gráfico de confianza
            st.plotly_chart(viz_data['confidence_chart'], use_container_width=True)
            
            # Tabla de datos
            st.markdown("#### 📈 Datos de Visualización")
            df_viz = pd.DataFrame(viz_data['sentiments_data'])
            st.dataframe(df_viz, use_container_width=True)
        
    except Exception as e:
        st.error(f"Error generando visualizaciones: {str(e)}")
        st.info("💡 Intenta regenerar el análisis")


def render_sentiment_by_source(result: SentimentAnalysisResult):
    """Renderizar análisis de sentimientos por fuente"""
    st.markdown("#### 🔗 Análisis por Fuente")
    
    # Agrupar sentimientos por fuente
    source_sentiments = {}
    for sentiment in result.sentiments:
        for citation in sentiment.citations:
            source = citation.get('source_file', 'unknown')
            if source not in source_sentiments:
                source_sentiments[source] = []
            source_sentiments[source].append(sentiment)
    
    if not source_sentiments:
        st.info("No hay datos de fuentes disponibles")
        return
    
    # Mostrar análisis por fuente
    for source, sentiments in source_sentiments.items():
        with st.expander(f"📄 {source} ({len(sentiments)} sentimientos)"):
            # Calcular estadísticas por fuente
            polarities = [s.polarity for s in sentiments]
            subjectivities = [s.subjectivity for s in sentiments]
            confidences = [s.confidence for s in sentiments]
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.metric("Polaridad Promedio", f"{np.mean(polarities):.3f}")
                st.metric("Subjetividad Promedio", f"{np.mean(subjectivities):.3f}")
            
            with col2:
                st.metric("Confianza Promedio", f"{np.mean(confidences):.3f}")
                st.metric("Total Sentimientos", len(sentiments))
            
            with col3:
                # Distribución de sentimientos por fuente
                sentiment_counts = {}
                for s in sentiments:
                    sentiment_counts[s.sentiment_label] = sentiment_counts.get(s.sentiment_label, 0) + 1
                
                st.markdown("**Distribución:**")
                for label, count in sentiment_counts.items():
                    st.markdown(f"• {label}: {count}")


def render_sentiment_export(result: SentimentAnalysisResult):
    """Renderizar opciones de exportación"""
    st.markdown("#### 💾 Exportar Análisis de Sentimientos")
    
    st.markdown("""
    <div style="background: #34495e; padding: 1rem; border-radius: 8px; border-left: 4px solid #e67e22; color: #ecf0f1; margin: 1rem 0;">
        <strong>ℹ️ Exportación Completa:</strong> Genera documentos profesionales con todos los sentimientos analizados,
        incluyendo visualizaciones, patrones emocionales y métricas de calidad para uso en tu investigación.
    </div>
    """, unsafe_allow_html=True)
    
    # Crear extractor para exportación
    from ...extractors.sentiment_extractor import SentimentExtractor
    extractor = SentimentExtractor(AnalysisConfig())
    
    # Opciones de exportación
    col1, col2 = st.columns(2)
    
    with col1:
        include_emotions = st.checkbox(
            "Incluir análisis de emociones",
            value=True,
            help="Incluye las emociones específicas identificadas por la IA",
            key="sentiment_include_emotions"
        )
    
    with col2:
        include_visualizations = st.checkbox(
            "Incluir datos de visualización",
            value=True,
            help="Incluye métricas de polaridad y subjetividad",
            key="sentiment_include_visualizations"
        )
    
    # Información del documento
    st.markdown("#### 📄 Información del Documento")
    
    col1, col2 = st.columns(2)
    
    with col1:
        document_title = st.text_input(
            "Título del documento",
            value="Análisis de Sentimientos",
            help="Título que aparecerá en el documento",
            key="sentiment_document_title"
        )
    
    with col2:
        author_name = st.text_input(
            "Autor",
            value="Investigador",
            help="Nombre del autor del análisis",
            key="sentiment_author_name"
        )
    
    # Botones de exportación
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("📄 Generar Documento Word", type="primary", use_container_width=True, key="sentiment_generate_word"):
            try:
                doc_content = _generate_sentiment_word_document(
                    result=result,
                    title=document_title,
                    author=author_name,
                    include_emotions=include_emotions,
                    include_visualizations=include_visualizations
                )
                
                st.download_button(
                    label="💾 Descargar Documento Word",
                    data=doc_content,
                    file_name=f"{document_title.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                st.success("✅ Documento Word generado correctamente")
                
            except Exception as e:
                st.error(f"❌ Error al generar documento Word: {str(e)}")
                st.exception(e)
    
    with col2:
        # Exportar datos
        export_data = extractor.export_sentiment_analysis(result, 'json')
        
        if export_data:
            import json
            json_data = json.dumps(export_data, indent=2, ensure_ascii=False)
            
            st.download_button(
                label="📥 Descargar Análisis (JSON)",
                data=json_data,
                file_name=f"sentiment_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                mime="application/json"
            )
    
    # Mostrar resumen de exportación
    if export_data:
        st.markdown("#### 📋 Resumen de Exportación")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**Metadatos:**")
            st.json(export_data['metadata'])
        
        with col2:
            st.markdown("**Estadísticas:**")
            st.metric("Total Analizado", len(export_data['sentiments']))
            st.metric("Polaridad Promedio", f"{export_data['metadata']['avg_polarity']:.3f}")
            st.metric("Subjetividad Promedio", f"{export_data['metadata']['avg_subjectivity']:.3f}")
        
        st.info("""
        **💡 Información sobre la exportación:**
        
        - **Word**: Documento profesional con sentimientos, emociones y métricas
        - **JSON**: Incluye todos los datos del análisis para procesamiento posterior
        - **Metadatos**: Información sobre el algoritmo y parámetros utilizados
        - **Emociones**: Análisis emocional específico identificado por la IA
        - **Estadísticas**: Métricas de polaridad y subjetividad
        """)
    
    # Botón para limpiar resultados
    if st.button("🗑️ Limpiar Resultados", type="secondary", key="sentiment_clear_results"):
        if 'sentiments_analyzed' in st.session_state:
            del st.session_state.sentiments_analyzed
        if 'sentiment_analysis_result' in st.session_state:
            del st.session_state.sentiment_analysis_result
        if 'sentiment_extractor' in st.session_state:
            del st.session_state.sentiment_extractor
        if 'sentiments_summary' in st.session_state:
            del st.session_state.sentiments_summary
        st.rerun()


def _generate_sentiment_word_document(
    result: SentimentAnalysisResult,
    title: str,
    author: str,
    include_emotions: bool = True,
    include_visualizations: bool = True
) -> bytes:
    """
    Generar documento Word con los sentimientos analizados
    
    Args:
        result: Resultado del análisis de sentimientos
        title: Título del documento
        author: Autor del documento
        include_emotions: Si incluir análisis de emociones
        include_visualizations: Si incluir métricas de visualización
        
    Returns:
        Contenido del documento Word como bytes
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        import io
    except ImportError:
        raise ImportError("python-docx es requerido para generar documentos Word. Instala con: pip install python-docx")
    
    # Crear documento
    doc = Document()
    
    # Configurar estilos
    title_style = doc.styles['Title']
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    
    heading_style = doc.styles['Heading 1']
    heading_style.font.size = Pt(14)
    heading_style.font.bold = True
    
    # Título del documento
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Información del documento
    doc.add_paragraph(f"Autor: {author}")
    doc.add_paragraph(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph("")
    
    # Resumen ejecutivo
    doc.add_heading("Resumen Ejecutivo", level=1)
    
    summary_text = f"""
    Este documento presenta un análisis cualitativo de sentimientos extraídos de documentos de investigación.
    
    Estadísticas del análisis:
    • Total de sentimientos analizados: {len(result.sentiments)}
    • Algoritmo utilizado: {result.algorithm_used}
    • Polaridad promedio: {result.average_polarity:.3f}
    • Subjetividad promedio: {result.average_subjectivity:.3f}
    """
    
    doc.add_paragraph(summary_text)
    doc.add_paragraph("")
    
    # Sentimientos principales
    doc.add_heading("Sentimientos Identificados", level=1)
    
    for i, sentiment in enumerate(result.sentiments, 1):
        # Título del sentimiento
        sentiment_title = f"{i}. {sentiment.sentiment_label.upper()}"
        doc.add_heading(sentiment_title, level=2)
        
        # Información básica
        info_text = f"""
        Polaridad: {sentiment.polarity:.3f}
        Subjetividad: {sentiment.subjectivity:.3f}
        Confianza: {sentiment.confidence:.3f}
        """
        
        doc.add_paragraph(info_text)
        
        # Texto analizado
        doc.add_heading("Texto Analizado", level=3)
        doc.add_paragraph(str(sentiment.text) if sentiment.text else "No disponible")
        
        # Contexto del sentimiento
        if sentiment.context:
            doc.add_heading("Contexto", level=3)
            doc.add_paragraph(str(sentiment.context) if sentiment.context else "No disponible")
        
        # Análisis de emociones
        if include_emotions and sentiment.metadata.get('emociones_especificas'):
            doc.add_heading("Emociones Específicas", level=3)
            emociones = sentiment.metadata['emociones_especificas']
            doc.add_paragraph(str(emociones) if emociones else "No disponible")
        
        # Intensidad emocional
        if include_emotions and sentiment.metadata.get('intensidad_emocional'):
            doc.add_heading("Intensidad Emocional", level=3)
            intensidad = sentiment.metadata['intensidad_emocional']
            if isinstance(intensidad, (int, float)):
                doc.add_paragraph(f"Intensidad: {intensidad:.2f}")
            else:
                doc.add_paragraph(str(intensidad))
        
        # Relevancia para la investigación
        if sentiment.metadata.get('relevancia_investigacion'):
            doc.add_heading("Relevancia para la Investigación", level=3)
            relevancia = sentiment.metadata['relevancia_investigacion']
            doc.add_paragraph(str(relevancia) if relevancia else "No disponible")
        
        # Citas y referencias
        if sentiment.citations:
            doc.add_heading("Referencias", level=3)
            
            for j, citation in enumerate(sentiment.citations[:3], 1):  # Máximo 3 citas
                if isinstance(citation, dict):
                    citation_text = f"""
                    Cita {j}:
                    Fuente: {citation.get('source_file', 'unknown')}
                    Contenido: {citation.get('content', '')[:200]}
                    """
                else:
                    citation_text = f"""
                    Cita {j}:
                    Fuente: {citation.source_file}
                    Contenido: {citation.get_full_context(max_chars=200)}
                    """
                doc.add_paragraph(citation_text)
        
        doc.add_paragraph("")  # Espacio entre sentimientos
    
    # Análisis estadístico
    if include_visualizations:
        doc.add_heading("Análisis Estadístico", level=1)
        
        # Estadísticas por tipo de sentimiento
        sentiment_types = [s.sentiment_label for s in result.sentiments]
        type_counts = {}
        for sent_type in sentiment_types:
            type_counts[sent_type] = type_counts.get(sent_type, 0) + 1
        
        doc.add_heading("Distribución por Tipo de Sentimiento", level=2)
        for sent_type, count in type_counts.items():
            doc.add_paragraph(f"• {sent_type}: {count} análisis")
        
        # Sentimientos más intensos
        doc.add_heading("Sentimientos Más Intensos", level=2)
        top_sentiments = sorted(result.sentiments, key=lambda x: abs(x.polarity), reverse=True)[:5]
        
        for i, sent in enumerate(top_sentiments, 1):
            doc.add_paragraph(f"{i}. {sent.sentiment_label} (Polaridad: {sent.polarity:.3f})")
    
    # Guardar en bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes.getvalue()
