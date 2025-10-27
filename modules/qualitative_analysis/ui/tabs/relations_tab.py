"""
Tab de Análisis de Relaciones
Interfaz de usuario para identificar y visualizar relaciones entre conceptos
"""

import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
import networkx as nx
from collections import defaultdict
from typing import List, Dict, Any, Optional
from datetime import datetime

from ...extractors.relation_extractor import RelationExtractor, RelationAnalysisResult
from ...core.config import AnalysisConfig
from ..components.educational import (
    show_methodology_box,
    show_interpretation_guide,
    show_citation_box,
    show_statistics_panel
)


def render_relations_tab(chunks: List[Dict[str, Any]], config: AnalysisConfig):
    """
    Renderizar tab de análisis de relaciones
    
    Args:
        chunks: Lista de chunks de texto para análisis
        config: Configuración del análisis
    """
    
    # Título con descripción
    st.markdown("""
    <div style="
        background: linear-gradient(90deg, #3498db 0%, #2980b9 100%);
        padding: 1.5rem;
        border-radius: 10px;
        margin-bottom: 2rem;
        text-align: center;
    ">
        <h1 style="color: white; margin: 0; font-size: 2rem;">🔗 Análisis de Relaciones</h1>
        <p style="color: white; margin: 0.5rem 0 0 0; font-size: 1.1rem; opacity: 0.9;">
            Identifica conexiones profundas entre conceptos con fundamentación completa
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # Mostrar metodología
    show_methodology_box(
        title="Enfoque Híbrido: Algoritmo + DeepSeek R1",
        description="""
        **¿Qué hace este análisis?**
        
        Este módulo identifica automáticamente las relaciones más significativas entre conceptos 
        utilizando un **enfoque híbrido** que combina algoritmos avanzados de análisis de redes 
        con inteligencia artificial para generar relaciones académicas profundas y explicaciones detalladas.
        """,
        steps=[
            "Preprocesamiento: Limpieza y normalización del texto",
            "Vectorización: Conversión a representación numérica (TF-IDF)",
            "Algoritmo de Análisis: Co-ocurrencias, Semántico o Causal para extracción inicial",
            "🤖 Refinamiento IA: DeepSeek R1 analiza y mejora las relaciones candidatas",
            "Generación de relaciones profundas: Crea conexiones académicas significativas",
            "Validación: Cálculo de fuerza mejorada y métricas de confianza",
            "Citación: Vinculación con fuentes originales",
            "Visualización: Redes interactivas y análisis de patrones"
        ],
        algorithm_info="Co-ocurrencias/Semántico/Causal + DeepSeek R1 para relaciones académicas profundas"
    )
    
    # Mostrar guía de interpretación
    show_interpretation_guide(
        what_it_means="""
        **¿Qué significan estos resultados?**
        
        Las relaciones identificadas son **conexiones académicas profundas** que emergen de tus documentos. 
        Con el refinamiento IA, cada relación captura patrones, dependencias y estructuras conceptuales 
        complejas que van más allá de simples co-ocurrencias, proporcionando comprensión matizada de 
        los sistemas conceptuales presentes en tu investigación.
        """,
        how_to_use=[
            "Identifica los patrones centrales de tu investigación",
            "Analiza las explicaciones generadas por la IA para cada relación",
            "Examina los patrones conceptuales para entender estructuras emergentes",
            "Estudia las redes de relaciones para entender sistemas complejos",
            "Usa las métricas de fuerza y confianza para validar la calidad",
            "Aprovecha las relaciones profundas para desarrollar marcos teóricos",
            "Exporta los resultados para fundamentar tu análisis cualitativo"
        ],
        limitations=[
            "Las relaciones generadas por IA requieren validación del investigador",
            "La calidad depende de la cantidad y diversidad del contenido analizado",
            "El LLM puede generar relaciones que sintetizan información no explícita",
            "Las explicaciones son interpretaciones que deben ser verificadas",
            "Requiere acceso a un modelo LLM (Ollama) para funcionar completamente"
        ]
    )
    
    # Verificar si hay chunks disponibles
    if not chunks:
        st.warning("⚠️ No hay documentos disponibles para análisis de relaciones")
        st.info("""
        **Para realizar análisis de relaciones:**
        1. Asegúrate de que hay documentos procesados en el cache RAG
        2. Los documentos deben contener suficientes conceptos relacionados
        3. Mínimo recomendado: 5 documentos o 1000 palabras
        """)
        return
    
    # Mostrar estadísticas de entrada
    st.markdown("#### 📊 Estadísticas de Entrada")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric(
            "📄 Documentos",
            len(set(chunk.get('metadata', {}).get('source_file', 'unknown') 
                   for chunk in chunks)),
            help="Número de documentos únicos"
        )
    
    with col2:
        st.metric(
            "📝 Chunks",
            len(chunks),
            help="Número total de fragmentos de texto"
        )
    
    with col3:
        st.metric(
            "📖 Palabras Estimadas",
            sum(len(chunk.get('content', '').split()) for chunk in chunks),
            help="Número aproximado de palabras"
        )
    
    with col4:
        st.metric(
            "📂 Fuentes Únicas",
            len(set(chunk.get('metadata', {}).get('source_file', 'unknown') 
                   for chunk in chunks)),
            help="Número de fuentes diferentes"
        )
    
    # Configuración del análisis (valores predefinidos en el código)
    st.markdown("#### ⚙️ Configuración del Análisis")
    
    st.info("""
    **ℹ️ Configuración Automática:** Los parámetros del análisis de relaciones están optimizados automáticamente.
    Si necesitas modificar algún parámetro específico, puedes hacerlo directamente en el código.
    """)
    
    st.divider()
    
    # Botones de análisis
    col1, col2, col3 = st.columns([2, 1, 2])
    
    with col1:
        if st.session_state.get('relations_analyzed', False):
            if st.button("🔄 Nuevo Análisis", type="secondary", use_container_width=True, help="Limpiar resultados y realizar nuevo análisis"):
                # Limpiar session state
                st.session_state.relations_analyzed = False
                st.session_state.relation_analysis_result = None
                st.session_state.relation_extractor = None
                st.session_state.relations_summary = None
                st.rerun()
    
    with col2:
        if not st.session_state.get('relations_analyzed', False):
            analyze_button = st.button(
                "🚀 Analizar Relaciones",
                type="primary",
                use_container_width=True,
                help="Iniciar análisis híbrido de relaciones"
            )
        else:
            analyze_button = False
    
    with col3:
        if st.session_state.get('relations_analyzed', False):
            st.success("✅ Análisis completado")
    
    # Realizar análisis solo si se presiona el botón Y no hay resultados previos
    if analyze_button and not st.session_state.get('relations_analyzed', False):
        with st.spinner("🔗 Analizando relaciones con enfoque híbrido..."):
            try:
                # Crear extractor
                extractor = RelationExtractor(config)
                
                # Ejecutar análisis híbrido
                result = extractor.analyze_relations_hybrid(chunks)
                
                # Guardar en session state
                st.session_state.relations_analyzed = True
                st.session_state.relation_analysis_result = result
                st.session_state.relation_extractor = extractor
                
                # Obtener resumen
                summary = extractor.get_relation_summary(result)
                st.session_state.relations_summary = summary
                
                # Forzar re-ejecución para mostrar resultados
                st.rerun()
                
            except Exception as e:
                st.error(f"❌ **Error en el análisis:** {str(e)}")
                st.exception(e)
                return
    
    # Mostrar resultados si ya están disponibles
    if st.session_state.get('relations_analyzed', False):
        # Obtener datos de session state
        result = st.session_state.relation_analysis_result
        extractor = st.session_state.relation_extractor
        summary = st.session_state.relations_summary
        
        # Mostrar resultados
        st.success(f"✅ **Análisis completado:** {len(result.relations)} relaciones identificadas")
        
        st.divider()
        
        # Panel de estadísticas
        show_statistics_panel(summary)
        
        st.divider()
        
        # Tabs para diferentes vistas
        tab1, tab2, tab3, tab4 = st.tabs([
            "🔗 Relaciones Principales",
            "🕸️ Red de Conceptos",
            "📊 Análisis Estadístico",
            "💾 Exportar"
        ])
        
        # Tab 1: Relaciones principales
        with tab1:
            render_relations_list(result)
        
        # Tab 2: Red de conceptos
        with tab2:
            render_network_graph(result)
        
        # Tab 3: Análisis estadístico
        with tab3:
            render_statistical_analysis(result)
        
        # Tab 4: Exportar
        with tab4:
            render_relation_export(result)


def render_relation_results(
    result: RelationAnalysisResult,
    config: Dict[str, Any],
    chunks: List[Dict[str, Any]]
):
    """
    Renderizar resultados del análisis de relaciones
    
    Args:
        result: Resultado del análisis
        config: Configuración utilizada
        chunks: Chunks originales para referencias
    """
    st.markdown("#### 📊 Resultados del Análisis de Relaciones")
    
    # Mostrar resumen del análisis
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Conceptos Analizados", len(result.concepts))
    
    with col2:
        st.metric("Relaciones Identificadas", result.total_relations)
    
    with col3:
        st.metric("Tipo de Análisis", result.analysis_type.capitalize())
    
    with col4:
        if result.relations:
            avg_strength = np.mean([r.strength for r in result.relations])
            st.metric("Fuerza Promedio", f"{avg_strength:.3f}")
        else:
            st.metric("Fuerza Promedio", "N/A")
    
    # Tabs para diferentes vistas de resultados
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "📋 Lista de Relaciones",
        "🕸️ Grafo de Red",
        "📊 Matriz de Relaciones",
        "📈 Análisis Estadístico",
        "💾 Exportar"
    ])
    
    with tab1:
        render_relations_list(result)
    
    with tab2:
        if config.get('show_network_graph', True):
            render_network_graph(result)
        else:
            st.info("Visualización de grafo deshabilitada en configuración")
    
    with tab3:
        if config.get('show_matrix', True):
            render_relation_matrix(result)
        else:
            st.info("Matriz de relaciones deshabilitada en configuración")
    
    with tab4:
        render_statistical_analysis(result)
    
    with tab5:
        render_relation_export(result)


def render_relations_list(result: RelationAnalysisResult):
    """Renderizar lista de relaciones identificadas"""
    st.markdown("#### 📋 Relaciones Identificadas")
    
    if not result.relations:
        st.info("No se identificaron relaciones con los parámetros seleccionados")
        return
    
    # Crear DataFrame para mostrar relaciones
    relations_data = []
    for i, relation in enumerate(result.relations):
        relations_data.append({
            'ID': i + 1,
            'Concepto Origen': relation.source_concept,
            'Concepto Destino': relation.target_concept,
            'Tipo': relation.relation_type,
            'Fuerza': f"{relation.strength:.3f}",
            'Confianza': f"{relation.confidence:.3f}",
            'Contexto': relation.context[:80] + "..." if len(relation.context) > 80 else relation.context
        })
    
    df = pd.DataFrame(relations_data)
    
    # Mostrar tabla con relaciones
    st.dataframe(df, use_container_width=True)
    
    # Mostrar detalles de cada relación
    st.markdown("#### 🔍 Detalles de Relaciones")
    
    # Filtros para la lista
    col1, col2, col3 = st.columns(3)
    
    with col1:
        filter_type = st.selectbox(
            "Filtrar por Tipo",
            ["Todos"] + list(set(r.relation_type for r in result.relations)),
            key="filter_relation_type"
        )
    
    with col2:
        sort_by = st.selectbox(
            "Ordenar por",
            ["Fuerza", "Confianza", "Tipo"],
            key="sort_relations_by"
        )
    
    with col3:
        show_count = st.slider(
            "Mostrar",
            min_value=1,
            max_value=max(5, min(50, len(result.relations))),
            value=min(10, len(result.relations)),
            key="show_relations_count"
        )
    
    # Aplicar filtros
    filtered_relations = result.relations
    if filter_type != "Todos":
        filtered_relations = [r for r in filtered_relations if r.relation_type == filter_type]
    
    # Ordenar
    if sort_by == "Fuerza":
        filtered_relations = sorted(filtered_relations, key=lambda x: x.strength, reverse=True)
    elif sort_by == "Confianza":
        filtered_relations = sorted(filtered_relations, key=lambda x: x.confidence, reverse=True)
    else:  # Tipo
        filtered_relations = sorted(filtered_relations, key=lambda x: x.relation_type)
    
    # Mostrar relaciones filtradas
    for i, relation in enumerate(filtered_relations[:show_count]):
        with st.expander(f"**Relación {i+1}: {relation.source_concept} ↔ {relation.target_concept}**"):
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("**Información de la Relación:**")
                st.markdown(f"• **Tipo:** {relation.relation_type}")
                st.markdown(f"• **Fuerza:** {relation.strength:.3f}")
                st.markdown(f"• **Confianza:** {relation.confidence:.3f}")
                
                st.markdown("**Contexto:**")
                st.info(relation.context)
            
            with col2:
                st.markdown("**Métricas Visuales:**")
                
                # Barra de progreso para fuerza
                st.progress(relation.strength, text=f"Fuerza: {relation.strength:.1%}")
                
                # Barra de progreso para confianza
                st.progress(relation.confidence, text=f"Confianza: {relation.confidence:.1%}")
                
                # Metadata adicional
                if relation.metadata:
                    st.markdown("**Metadata:**")
                    st.json(relation.metadata)
            
            # Mostrar citaciones si existen
            if relation.citations:
                st.markdown("---")
                st.markdown("📚 **Citaciones:**")
                
                for j, citation in enumerate(relation.citations[:3], 1):
                    with st.container():
                        st.markdown(f"**Cita {j}:**")
                        st.markdown(f"📄 **Fuente:** {citation.get('source_file', 'unknown')}")
                        st.markdown(f"📝 **Contenido:** {citation.get('content', '')[:200]}...")
                        if j < len(relation.citations[:3]):
                            st.markdown("---")


def render_network_graph(result: RelationAnalysisResult):
    """Renderizar grafo de red interactivo"""
    st.markdown("#### 🕸️ Grafo de Red de Conceptos")
    
    if not result.relations:
        st.info("No hay relaciones para visualizar")
        return
    
    try:
        # Crear grafo de NetworkX
        G = nx.Graph()
        
        # Agregar nodos
        for concept in result.concepts:
            G.add_node(concept)
        
        # Agregar aristas
        for relation in result.relations:
            G.add_edge(
                relation.source_concept,
                relation.target_concept,
                weight=relation.strength,
                relation_type=relation.relation_type
            )
        
        # Calcular layout
        pos = nx.spring_layout(G, k=2, iterations=50)
        
        # Crear visualización con Plotly
        edge_trace = []
        for edge in G.edges():
            x0, y0 = pos[edge[0]]
            x1, y1 = pos[edge[1]]
            
            # Obtener información de la arista
            edge_info = G[edge[0]][edge[1]]
            weight = edge_info.get('weight', 0.5)
            
            edge_trace.append(
                go.Scatter(
                    x=[x0, x1, None],
                    y=[y0, y1, None],
                    mode='lines',
                    line=dict(width=weight * 5, color='rgba(125, 125, 125, 0.5)'),
                    hoverinfo='none'
                )
            )
        
        # Crear trazas de nodos
        node_x = []
        node_y = []
        node_text = []
        node_size = []
        
        for node in G.nodes():
            x, y = pos[node]
            node_x.append(x)
            node_y.append(y)
            
            # Calcular grado del nodo (número de conexiones)
            degree = G.degree(node)
            node_size.append(20 + degree * 5)
            node_text.append(f"{node}<br>Conexiones: {degree}")
        
        node_trace = go.Scatter(
            x=node_x,
            y=node_y,
            mode='markers+text',
            text=[node for node in G.nodes()],
            textposition="top center",
            marker=dict(
                size=node_size,
                color='lightblue',
                line=dict(width=2, color='darkblue')
            ),
            hovertext=node_text,
            hoverinfo='text'
        )
        
        # Crear figura
        fig = go.Figure(
            data=edge_trace + [node_trace],
            layout=go.Layout(
                title="Red de Relaciones entre Conceptos",
                showlegend=False,
                hovermode='closest',
                margin=dict(b=0, l=0, r=0, t=40),
                xaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
                yaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
                height=600
            )
        )
        
        st.plotly_chart(fig, use_container_width=True)
        
        # Estadísticas del grafo
        st.markdown("#### 📊 Estadísticas del Grafo")
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("Nodos", G.number_of_nodes())
        
        with col2:
            st.metric("Aristas", G.number_of_edges())
        
        with col3:
            density = nx.density(G)
            st.metric("Densidad", f"{density:.3f}")
        
        with col4:
            if G.number_of_nodes() > 0:
                avg_degree = sum(dict(G.degree()).values()) / G.number_of_nodes()
                st.metric("Grado Promedio", f"{avg_degree:.2f}")
        
    except Exception as e:
        st.error(f"Error generando grafo: {str(e)}")
        st.info("💡 Intenta con menos conceptos o relaciones más fuertes")


def render_relation_matrix(result: RelationAnalysisResult):
    """Renderizar matriz de relaciones"""
    st.markdown("#### 📊 Matriz de Relaciones")
    
    if not result.relations:
        st.info("No hay relaciones para mostrar en matriz")
        return
    
    try:
        # Crear matriz de relaciones
        matrix_size = len(result.concepts)
        relation_matrix = np.zeros((matrix_size, matrix_size))
        
        # Llenar matriz con fuerzas de relaciones
        for relation in result.relations:
            try:
                idx1 = result.concepts.index(relation.source_concept)
                idx2 = result.concepts.index(relation.target_concept)
                relation_matrix[idx1, idx2] = relation.strength
                relation_matrix[idx2, idx1] = relation.strength
            except ValueError:
                continue
        
        # Crear heatmap con Plotly
        fig = go.Figure(data=go.Heatmap(
            z=relation_matrix,
            x=result.concepts,
            y=result.concepts,
            colorscale='Viridis',
            hovertemplate='%{y} ↔ %{x}<br>Fuerza: %{z:.3f}<extra></extra>'
        ))
        
        fig.update_layout(
            title="Matriz de Fuerza de Relaciones",
            xaxis_title="Conceptos",
            yaxis_title="Conceptos",
            height=600
        )
        
        st.plotly_chart(fig, use_container_width=True)
        
        # Tabla de relaciones más fuertes
        st.markdown("#### 🔝 Top Relaciones Más Fuertes")
        
        # Ordenar relaciones por fuerza
        top_relations = sorted(result.relations, key=lambda x: x.strength, reverse=True)[:10]
        
        top_data = []
        for i, rel in enumerate(top_relations, 1):
            top_data.append({
                'Ranking': i,
                'Concepto 1': rel.source_concept,
                'Concepto 2': rel.target_concept,
                'Tipo': rel.relation_type,
                'Fuerza': f"{rel.strength:.3f}"
            })
        
        df_top = pd.DataFrame(top_data)
        st.dataframe(df_top, use_container_width=True)
        
    except Exception as e:
        st.error(f"Error generando matriz: {str(e)}")


def render_statistical_analysis(result: RelationAnalysisResult):
    """Renderizar análisis estadístico de relaciones"""
    st.markdown("#### 📈 Análisis Estadístico")
    
    if not result.relations:
        st.info("No hay datos para análisis estadístico")
        return
    
    # Estadísticas por tipo de relación
    st.markdown("##### 📊 Distribución por Tipo de Relación")
    
    relation_types = [r.relation_type for r in result.relations]
    type_counts = pd.Series(relation_types).value_counts()
    
    fig_types = px.pie(
        values=type_counts.values,
        names=type_counts.index,
        title="Distribución de Tipos de Relación"
    )
    st.plotly_chart(fig_types, use_container_width=True)
    
    # Distribución de fuerza de relaciones
    st.markdown("##### 📊 Distribución de Fuerza de Relaciones")
    
    strengths = [r.strength for r in result.relations]
    
    fig_strength = px.histogram(
        x=strengths,
        nbins=20,
        title="Distribución de Fuerza de Relaciones",
        labels={'x': 'Fuerza', 'y': 'Frecuencia'}
    )
    st.plotly_chart(fig_strength, use_container_width=True)
    
    # Estadísticas descriptivas
    st.markdown("##### 📊 Estadísticas Descriptivas")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric("Fuerza Media", f"{np.mean(strengths):.3f}")
        st.metric("Fuerza Mediana", f"{np.median(strengths):.3f}")
    
    with col2:
        st.metric("Fuerza Máxima", f"{np.max(strengths):.3f}")
        st.metric("Fuerza Mínima", f"{np.min(strengths):.3f}")
    
    with col3:
        st.metric("Desviación Estándar", f"{np.std(strengths):.3f}")
        st.metric("Coeficiente de Variación", f"{(np.std(strengths) / np.mean(strengths)):.3f}")
    
    # Conceptos más conectados
    st.markdown("##### 🌟 Conceptos Más Conectados")
    
    concept_connections = defaultdict(int)
    for relation in result.relations:
        concept_connections[relation.source_concept] += 1
        concept_connections[relation.target_concept] += 1
    
    top_concepts = sorted(concept_connections.items(), key=lambda x: x[1], reverse=True)[:10]
    
    top_concepts_data = []
    for i, (concept, connections) in enumerate(top_concepts, 1):
        top_concepts_data.append({
            'Ranking': i,
            'Concepto': concept,
            'Conexiones': connections
        })
    
    df_concepts = pd.DataFrame(top_concepts_data)
    st.dataframe(df_concepts, use_container_width=True)


def render_relation_export(result: RelationAnalysisResult):
    """Renderizar opciones de exportación"""
    st.markdown("#### 💾 Exportar Análisis de Relaciones")
    
    st.markdown("""
    <div style="background: #34495e; padding: 1rem; border-radius: 8px; border-left: 4px solid #e67e22; color: #ecf0f1; margin: 1rem 0;">
        <strong>ℹ️ Exportación Completa:</strong> Genera documentos profesionales con todas las relaciones analizadas,
        incluyendo visualizaciones, patrones conceptuales y métricas de calidad para uso en tu investigación.
    </div>
    """, unsafe_allow_html=True)
    
    # Crear extractor para exportación
    from ...extractors.relation_extractor import RelationExtractor
    extractor = RelationExtractor(AnalysisConfig())
    
    # Opciones de exportación
    col1, col2 = st.columns(2)
    
    with col1:
        include_patterns = st.checkbox(
            "Incluir patrones conceptuales",
            value=True,
            help="Incluye los patrones conceptuales identificados por la IA"
        )
    
    with col2:
        include_visualizations = st.checkbox(
            "Incluir datos de visualización",
            value=True,
            help="Incluye métricas de fuerza y confianza"
        )
    
    # Información del documento
    st.markdown("#### 📄 Información del Documento")
    
    col1, col2 = st.columns(2)
    
    with col1:
        document_title = st.text_input(
            "Título del documento",
            value="Análisis de Relaciones",
            help="Título que aparecerá en el documento"
        )
    
    with col2:
        author_name = st.text_input(
            "Autor",
            value="Investigador",
            help="Nombre del autor del análisis"
        )
    
    # Botones de exportación
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("📄 Generar Documento Word", type="primary", use_container_width=True):
            try:
                doc_content = _generate_relation_word_document(
                    result=result,
                    title=document_title,
                    author=author_name,
                    include_patterns=include_patterns,
                    include_visualizations=include_visualizations
                )
                
                st.download_button(
                    label="💾 Descargar Documento Word",
                    data=doc_content,
                    file_name=f"{document_title.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                st.success("✅ Documento Word generado correctamente")
                
            except Exception as e:
                st.error(f"❌ Error al generar documento Word: {str(e)}")
                st.exception(e)
    
    with col2:
        # Exportar datos
        export_data = extractor.export_relations(result, 'json')
        
        if export_data:
            import json
            json_data = json.dumps(export_data, indent=2, ensure_ascii=False)
            
            st.download_button(
                label="📥 Descargar Análisis (JSON)",
                data=json_data,
                file_name=f"relation_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                mime="application/json"
            )
    
    # Mostrar resumen de exportación
    if export_data:
        st.markdown("#### 📋 Resumen de Exportación")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**Metadatos:**")
            st.json(export_data['metadata'])
        
        with col2:
            st.markdown("**Estadísticas:**")
            st.metric("Total Relaciones", len(export_data['relations']))
            st.metric("Total Conceptos", len(export_data['concepts']))
            st.metric("Tipo de Análisis", export_data['metadata']['analysis_type'])
        
        st.info("""
        **💡 Información sobre la exportación:**
        
        - **Word**: Documento profesional con relaciones, patrones y métricas
        - **JSON**: Incluye todos los datos del análisis para procesamiento posterior
        - **Metadatos**: Información sobre el algoritmo y parámetros utilizados
        - **Patrones**: Patrones conceptuales identificados por la IA
        - **Estadísticas**: Métricas de calidad y confianza de las relaciones
        """)
    
    # Botón para limpiar resultados
    if st.button("🗑️ Limpiar Resultados", type="secondary"):
        if 'relations_analyzed' in st.session_state:
            del st.session_state.relations_analyzed
        if 'relation_analysis_result' in st.session_state:
            del st.session_state.relation_analysis_result
        if 'relation_extractor' in st.session_state:
            del st.session_state.relation_extractor
        if 'relations_summary' in st.session_state:
            del st.session_state.relations_summary
        st.rerun()


def _generate_relation_word_document(
    result: RelationAnalysisResult,
    title: str,
    author: str,
    include_patterns: bool = True,
    include_visualizations: bool = True
) -> bytes:
    """
    Generar documento Word con las relaciones analizadas
    
    Args:
        result: Resultado del análisis de relaciones
        title: Título del documento
        author: Autor del documento
        include_patterns: Si incluir patrones conceptuales
        include_visualizations: Si incluir métricas de visualización
        
    Returns:
        Contenido del documento Word como bytes
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        import io
    except ImportError:
        raise ImportError("python-docx es requerido para generar documentos Word. Instala con: pip install python-docx")
    
    # Crear documento
    doc = Document()
    
    # Configurar estilos
    title_style = doc.styles['Title']
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    
    heading_style = doc.styles['Heading 1']
    heading_style.font.size = Pt(14)
    heading_style.font.bold = True
    
    # Título del documento
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Información del documento
    doc.add_paragraph(f"Autor: {author}")
    doc.add_paragraph(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph("")
    
    # Resumen ejecutivo
    doc.add_heading("Resumen Ejecutivo", level=1)
    
    summary_text = f"""
    Este documento presenta un análisis cualitativo de relaciones entre conceptos extraídos de documentos de investigación.
    
    Estadísticas del análisis:
    • Total de relaciones identificadas: {len(result.relations)}
    • Conceptos analizados: {len(result.concepts)}
    • Tipo de análisis: {result.analysis_type}
    • Fuerza promedio: {sum(r.strength for r in result.relations) / len(result.relations):.3f}
    • Confianza promedio: {sum(r.confidence for r in result.relations) / len(result.relations):.3f}
    """
    
    doc.add_paragraph(summary_text)
    doc.add_paragraph("")
    
    # Relaciones principales
    doc.add_heading("Relaciones Identificadas", level=1)
    
    for i, relation in enumerate(result.relations, 1):
        # Título de la relación
        relation_title = f"{i}. {relation.source_concept} ↔ {relation.target_concept}"
        doc.add_heading(relation_title, level=2)
        
        # Información básica
        info_text = f"""
        Tipo de relación: {relation.relation_type}
        Fuerza: {relation.strength:.3f}
        Confianza: {relation.confidence:.3f}
        """
        
        doc.add_paragraph(info_text)
        
        # Explicación de la relación
        if relation.context:
            doc.add_heading("Explicación", level=3)
            doc.add_paragraph(relation.context)
        
        # Patrones conceptuales
        if include_patterns and relation.metadata.get('pattern_conceptual'):
            doc.add_heading("Patrón Conceptual", level=3)
            doc.add_paragraph(relation.metadata['pattern_conceptual'])
        
        # Relevancia para la investigación
        if relation.metadata.get('relevancia_investigacion'):
            doc.add_heading("Relevancia para la Investigación", level=3)
            doc.add_paragraph(relation.metadata['relevancia_investigacion'])
        
        # Citas y referencias
        if relation.citations:
            doc.add_heading("Referencias", level=3)
            
            for j, citation in enumerate(relation.citations[:3], 1):  # Máximo 3 citas
                if isinstance(citation, dict):
                    citation_text = f"""
                    Cita {j}:
                    Fuente: {citation.get('source_file', 'unknown')}
                    Contenido: {citation.get('content', '')[:200]}
                    """
                else:
                    citation_text = f"""
                    Cita {j}:
                    Fuente: {citation.source_file}
                    Contenido: {citation.get_full_context(max_chars=200)}
                    """
                doc.add_paragraph(citation_text)
        
        doc.add_paragraph("")  # Espacio entre relaciones
    
    # Análisis estadístico
    if include_visualizations:
        doc.add_heading("Análisis Estadístico", level=1)
        
        # Estadísticas por tipo de relación
        relation_types = [r.relation_type for r in result.relations]
        type_counts = {}
        for rel_type in relation_types:
            type_counts[rel_type] = type_counts.get(rel_type, 0) + 1
        
        doc.add_heading("Distribución por Tipo de Relación", level=2)
        for rel_type, count in type_counts.items():
            doc.add_paragraph(f"• {rel_type}: {count} relaciones")
        
        # Relaciones más fuertes
        doc.add_heading("Relaciones Más Fuertes", level=2)
        top_relations = sorted(result.relations, key=lambda x: x.strength, reverse=True)[:5]
        
        for i, rel in enumerate(top_relations, 1):
            doc.add_paragraph(f"{i}. {rel.source_concept} ↔ {rel.target_concept} (Fuerza: {rel.strength:.3f})")
    
    # Guardar en bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes.getvalue()

