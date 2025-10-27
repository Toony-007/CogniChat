"""
Tab de Extracción de Conceptos Clave
Primer sub-módulo del sistema de análisis cualitativo
"""

import streamlit as st
from typing import List, Dict, Any
import json
from datetime import datetime

from ...extractors.concept_extractor import ConceptExtractor, ExtractedConcept
from ...core.config import AnalysisConfig
from ..components.educational import (
    show_methodology_box,
    show_interpretation_guide,
    show_concept_card,
    show_statistics_panel,
    show_help_sidebar
)


def render_concepts_tab(chunks: List[Dict[str, Any]], config: AnalysisConfig):
    """
    Renderizar tab de extracción de conceptos clave
    
    Args:
        chunks: Lista de chunks de documentos
        config: Configuración del análisis
    """
    
    # Título con descripción
    st.markdown("""
    <div style="
        background: linear-gradient(90deg, #00FF99 0%, #00CC7A 100%);
        padding: 1.5rem;
        border-radius: 10px;
        margin-bottom: 2rem;
        text-align: center;
    ">
        <h1 style="color: white; margin: 0; font-size: 2rem;">🔍 Extracción de Conceptos Clave</h1>
        <p style="color: white; margin: 0.5rem 0 0 0; font-size: 1.1rem; opacity: 0.9;">
            Identifica los conceptos más importantes de tus documentos con fundamentación completa
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # Mostrar ayuda en sidebar
    show_help_sidebar()
    
    # Mostrar metodología si está habilitado
    if config.show_methodology:
        methodology_title = "Enfoque Híbrido: TF-IDF + DeepSeek R1" if config.enable_llm_refinement else "Extracción Inteligente con TF-IDF"
        
        methodology_description = """
            Este análisis identifica los **conceptos más relevantes** de tus documentos
            usando un **enfoque híbrido** que combina:
            
            - **TF-IDF** (Term Frequency-Inverse Document Frequency) para extracción inicial
            - **DeepSeek R1** para refinamiento y mejora de conceptos
            
            El proceso identifica términos que son:
            - **Frecuentes** en documentos específicos (TF - Term Frequency)
            - **Raros** en el corpus general (IDF - Inverse Document Frequency)
            - **Refinados** por IA para generar conceptos académicos profundos
            """ if config.enable_llm_refinement else """
            Este análisis identifica los **conceptos más relevantes** de tus documentos
            usando **TF-IDF** (Term Frequency-Inverse Document Frequency), una técnica
            estándar en análisis de textos que identifica términos que son:
            
            - **Frecuentes** en documentos específicos (TF - Term Frequency)
            - **Raros** en el corpus general (IDF - Inverse Document Frequency)
            
            Esto permite encontrar conceptos que son **importantes y distintivos**.
            """
        
        methodology_steps = [
            "**Preprocesamiento**: Limpia el texto, elimina palabras vacías (stopwords) y normaliza",
            "**Vectorización**: Convierte el texto en representación matemática usando TF-IDF",
            "**Extracción inicial**: Identifica términos individuales y frases completas (n-gramas)",
            "**Puntuación**: Calcula un score de relevancia para cada concepto (0.0 a 1.0)",
            "**🤖 Refinamiento IA**: DeepSeek R1 analiza y mejora los conceptos candidatos",
            "**Generación de conceptos profundos**: Crea conceptos académicos significativos",
            "**Citación**: Identifica y registra todas las fuentes donde aparece cada concepto",
            "**Análisis de relaciones**: Identifica conceptos que aparecen juntos frecuentemente",
            "**Ordenamiento**: Presenta los conceptos ordenados por relevancia"
        ] if config.enable_llm_refinement else [
            "**Preprocesamiento**: Limpia el texto, elimina palabras vacías (stopwords) y normaliza",
            "**Vectorización**: Convierte el texto en representación matemática usando TF-IDF",
            "**Extracción**: Identifica términos individuales y frases completas (n-gramas)",
            "**Puntuación**: Calcula un score de relevancia para cada concepto (0.0 a 1.0)",
            "**Citación**: Identifica y registra todas las fuentes donde aparece cada concepto",
            "**Análisis de relaciones**: Identifica conceptos que aparecen juntos frecuentemente",
            "**Ordenamiento**: Presenta los conceptos ordenados por relevancia"
        ]
        
        algorithm_info = "TF-IDF + DeepSeek R1 para conceptos académicos profundos" if config.enable_llm_refinement else "TF-IDF con detección de n-gramas (frases de 1-3 palabras)"
        
        show_methodology_box(
            title=methodology_title,
            description=methodology_description,
            steps=methodology_steps,
            algorithm_info=algorithm_info
        )
    
    # Mostrar guía de interpretación si está habilitado
    if config.show_interpretation_guide:
        what_it_means = """
            Los **conceptos extraídos** representan las ideas centrales y términos técnicos
            más importantes en tus documentos. El **score de relevancia** indica qué tan
            característico es cada concepto del contenido analizado.
            
            Los conceptos con mayor score son aquellos que:
            - Aparecen frecuentemente en tus documentos
            - Son específicos de tu contenido (no son palabras comunes)
            - Están bien distribuidos en las fuentes
            """ if not config.enable_llm_refinement else """
            Los **conceptos extraídos** son ideas académicas profundas y significativas
            que emergen de tus documentos. El **score de relevancia** indica qué tan
            característico es cada concepto del contenido analizado.
            
            Con el refinamiento IA, los conceptos son:
            - **Académicos y profundos**: Capturan procesos, relaciones y fenómenos complejos
            - **Específicos al contexto**: Explican "cómo" y "por qué", no solo "qué"
            - **Fundamentados**: Incluyen explicaciones y categorización automática
            - **Significativos**: Eliminan palabras genéricas y letras sueltas
            """
        
        how_to_use = [
            "Identifica los temas principales de tu corpus documental",
            "Verifica que los conceptos extraídos coinciden con tu comprensión del contenido",
            "Usa las citas para volver a las fuentes originales y profundizar",
            "Examina los conceptos relacionados para entender conexiones temáticas",
            "Compara la distribución de conceptos entre diferentes fuentes",
            "Usa estos conceptos como base para análisis posteriores (temas, relaciones, etc.)"
        ] if not config.enable_llm_refinement else [
            "Identifica los fenómenos centrales de tu investigación",
            "Analiza las explicaciones generadas por la IA para cada concepto",
            "Examina las categorías asignadas (metodología, teoría, resultado, etc.)",
            "Usa las citas para volver a las fuentes originales y validar",
            "Estudia las relaciones entre conceptos para entender patrones complejos",
            "Aprovecha los conceptos profundos para desarrollar marcos teóricos",
            "Utiliza las explicaciones para fundamentar tu análisis cualitativo"
        ]
        
        limitations = [
            "El sistema identifica palabras y frases frecuentes, pero no comprende significado",
            "Conceptos muy específicos con baja frecuencia pueden no aparecer",
            "La relevancia es estadística, no semántica (requiere validación del investigador)",
            "Términos polisémicos (múltiples significados) no se distinguen automáticamente",
            "La calidad depende de la calidad y cantidad de documentos analizados"
        ] if not config.enable_llm_refinement else [
            "Los conceptos generados por IA requieren validación del investigador",
            "La calidad depende de la calidad y cantidad de documentos analizados",
            "El LLM puede generar conceptos que no están directamente en el texto",
            "Las explicaciones son interpretaciones que deben ser verificadas",
            "La categorización automática puede no ser perfecta",
            "Requiere acceso a un modelo LLM (Ollama) para funcionar"
        ]
        
        show_interpretation_guide(
            what_it_means=what_it_means,
            how_to_use=how_to_use,
            limitations=limitations
        )
    
    # Verificar que hay documentos
    if not chunks:
        st.warning("""
        ⚠️ **No hay documentos para analizar**
        
        Por favor, ve a la pestaña **"📄 Gestión de Documentos"** para:
        1. Subir documentos (PDF, DOCX, TXT, etc.)
        2. Luego ve a **"🧠 Procesamiento RAG"** para procesar los documentos
        3. Vuelve aquí para realizar el análisis cualitativo
        """)
        return
    
    # Información sobre los documentos
    st.markdown("### 📂 Documentos Disponibles")
    
    sources = list(set(chunk['metadata'].get('source_file', 'unknown') for chunk in chunks))
    total_content = sum(len(chunk.get('content', '')) for chunk in chunks)
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("📄 Documentos", len(sources))
    with col2:
        st.metric("📝 Chunks", len(chunks))
    with col3:
        st.metric("📊 Caracteres", f"{total_content:,}")
    
    with st.expander("📋 Ver lista de fuentes"):
        for i, source in enumerate(sources, 1):
            st.markdown(f"{i}. `{source}`")
    
    st.divider()
    
    # Configuración del análisis (valores predefinidos en el código)
    st.markdown("### ⚙️ Configuración del Análisis")
    
    st.info("""
    **ℹ️ Configuración Automática:** Los parámetros del análisis están optimizados automáticamente.
    Si necesitas modificar algún parámetro específico, puedes hacerlo directamente en el código.
    """)
    
    st.divider()
    
    # Configuración de categorías iniciales
    st.markdown("#### 🏷️ Categorías Iniciales (Opcional)")
    
    use_custom_categories = st.checkbox(
        "Usar categorías iniciales personalizadas",
        value=config.use_custom_categories,
        help="Define categorías específicas para dirigir la extracción de conceptos",
        key="concepts_use_custom_categories"
    )
    
    if use_custom_categories:
        st.markdown("**Define las categorías que necesitas identificar en tu investigación:**")
        
        # Inicializar categorías en session state si no existen
        if 'custom_categories' not in st.session_state:
            st.session_state.custom_categories = config.custom_categories.copy()
        
        # Mostrar categorías existentes
        categories_to_remove = []
        for i, (category, definition) in enumerate(st.session_state.custom_categories.items()):
            with st.expander(f"📂 {category}", expanded=True):
                col1, col2 = st.columns([3, 1])
                
                with col1:
                    new_category = st.text_input(
                        f"Categoría {i+1}",
                        value=category,
                        key=f"concepts_category_{i}",
                        help="Nombre de la categoría (ej: Metodología de investigación)"
                    )
                    
                    new_definition = st.text_area(
                        f"Definición {i+1}",
                        value=definition,
                        key=f"concepts_definition_{i}",
                        height=100,
                        help="Descripción detallada de qué incluye esta categoría"
                    )
                
                with col2:
                    if st.button("🗑️", key=f"concepts_remove_{i}", help="Eliminar categoría"):
                        categories_to_remove.append(category)
                
                # Actualizar categoría si cambió
                if new_category != category or new_definition != definition:
                    if new_category and new_definition:
                        # Eliminar la antigua
                        if category in st.session_state.custom_categories:
                            del st.session_state.custom_categories[category]
                        # Agregar la nueva
                        st.session_state.custom_categories[new_category] = new_definition
                        st.rerun()
        
        # Eliminar categorías marcadas para eliminación
        for category in categories_to_remove:
            if category in st.session_state.custom_categories:
                del st.session_state.custom_categories[category]
                st.rerun()
        
        # Botón para agregar nueva categoría
        if st.button("➕ Agregar Nueva Categoría", type="secondary", key="concepts_add_category"):
            new_key = f"new_category_{len(st.session_state.custom_categories)}"
            st.session_state.custom_categories[new_key] = ""
            st.rerun()
        
        # Mostrar resumen
        if st.session_state.custom_categories:
            st.success(f"✅ {len(st.session_state.custom_categories)} categorías definidas")
            with st.expander("📋 Ver resumen de categorías"):
                for category, definition in st.session_state.custom_categories.items():
                    st.markdown(f"**{category}:** {definition[:100]}{'...' if len(definition) > 100 else ''}")
        else:
            st.info("💡 Agrega al menos una categoría para dirigir la extracción de conceptos")
    
    st.divider()
    
    # Actualizar configuración (solo categorías personalizadas)
    config.use_custom_categories = use_custom_categories
    
    # Actualizar categorías personalizadas si están habilitadas
    if use_custom_categories and 'custom_categories' in st.session_state:
        config.custom_categories = st.session_state.custom_categories.copy()
    
    st.divider()
    
    # Botones de análisis
    col1, col2, col3 = st.columns([2, 1, 2])
    
    with col1:
        if st.session_state.get('concepts_extracted', False):
            if st.button("🔄 Nuevo Análisis", type="secondary", use_container_width=True, help="Limpiar resultados y realizar nuevo análisis", key="concepts_new_analysis"):
                # Limpiar session state
                st.session_state.concepts_extracted = False
                st.session_state.extracted_concepts = None
                st.session_state.concept_extractor = None
                st.session_state.concepts_summary = None
                st.rerun()
    
    with col2:
        if not st.session_state.get('concepts_extracted', False):
            analyze_button = st.button(
                "🚀 Extraer Conceptos",
                type="primary",
                use_container_width=True,
                help="Iniciar análisis de extracción de conceptos",
                key="concepts_analyze"
            )
        else:
            analyze_button = False
    
    with col3:
        if st.session_state.get('concepts_extracted', False):
            st.success("✅ Análisis completado")
    
    # Realizar análisis solo si se presiona el botón Y no hay resultados previos
    if analyze_button and not st.session_state.get('concepts_extracted', False):
        
        # Mostrar spinner durante análisis
        with st.spinner("🔍 Analizando documentos y extrayendo conceptos..."):
            try:
                # Crear extractor
                extractor = ConceptExtractor(config)
                
                # Extraer conceptos
                concepts = extractor.extract_concepts(chunks, method='tfidf')
                
                # Guardar en session state
                st.session_state.concepts_extracted = True
                st.session_state.extracted_concepts = concepts
                st.session_state.concept_extractor = extractor
                
                # Obtener resumen
                summary = extractor.get_concept_summary(concepts)
                st.session_state.concepts_summary = summary
                
                # Forzar re-ejecución para mostrar resultados
                st.rerun()
                
            except Exception as e:
                st.error(f"❌ **Error en el análisis:** {str(e)}")
                st.exception(e)
                return
    
    # Mostrar resultados si ya están disponibles
    if st.session_state.get('concepts_extracted', False):
        # Obtener datos de session state
        concepts = st.session_state.extracted_concepts
        extractor = st.session_state.concept_extractor
        summary = st.session_state.concepts_summary
        
        # Mostrar resultados
        st.success(f"✅ **Análisis completado:** {len(concepts)} conceptos extraídos")
        
        st.divider()
        
        # Panel de estadísticas
        show_statistics_panel(summary)
        
        st.divider()
        
        # Tabs para diferentes vistas
        tab1, tab2, tab3, tab4 = st.tabs([
            "🎯 Conceptos Principales",
            "🔗 Relaciones",
            "📚 Bibliografía",
            "💾 Exportar"
        ])
        
        # Tab 1: Conceptos principales
        with tab1:
            st.markdown("### 🎯 Conceptos Clave Identificados")
            
            st.markdown("""
            <div style="background: #34495e; padding: 1rem; border-radius: 8px; border-left: 4px solid #00FF99; color: #ecf0f1; margin: 1rem 0;">
                <strong>ℹ️ Interpretación:</strong> Los conceptos están ordenados por relevancia. 
                Los primeros son los más característicos de tu corpus documental.
                Cada concepto incluye <strong>citas a las fuentes originales</strong> para fundamentación.
            </div>
            """, unsafe_allow_html=True)
            
            # Filtros
            col1, col2 = st.columns([3, 1])
            with col1:
                filter_text = st.text_input(
                    "🔍 Filtrar conceptos",
                    placeholder="Escribe para buscar...",
                    help="Filtra los conceptos por texto",
                    key="concepts_filter_text"
                )
            with col2:
                sort_by = st.selectbox(
                    "Ordenar por",
                    ["Relevancia", "Frecuencia", "Nombre"],
                    help="Criterio de ordenamiento",
                    key="concepts_sort_by"
                )
            
            # Aplicar filtros
            filtered_concepts = concepts
            if filter_text:
                filtered_concepts = [
                    c for c in concepts
                    if filter_text.lower() in c.concept.lower()
                ]
            
            # Ordenar
            if sort_by == "Frecuencia":
                filtered_concepts = sorted(filtered_concepts, key=lambda c: c.frequency, reverse=True)
            elif sort_by == "Nombre":
                filtered_concepts = sorted(filtered_concepts, key=lambda c: c.concept)
            
            # Paginación
            concepts_per_page = 10
            total_pages = (len(filtered_concepts) + concepts_per_page - 1) // concepts_per_page
            
            if total_pages > 1:
                col1, col2, col3 = st.columns([2, 1, 2])
                with col2:
                    page = st.selectbox(
                        "Página",
                        range(1, total_pages + 1),
                        format_func=lambda x: f"Página {x} de {total_pages}",
                        key="concepts_page_selector"
                    )
            else:
                page = 1
            
            # Calcular rango
            start_idx = (page - 1) * concepts_per_page
            end_idx = min(start_idx + concepts_per_page, len(filtered_concepts))
            
            # Mostrar conceptos
            st.markdown(f"**Mostrando {start_idx + 1}-{end_idx} de {len(filtered_concepts)} conceptos**")
            
            for i, concept in enumerate(filtered_concepts[start_idx:end_idx], start=start_idx + 1):
                show_concept_card(
                    concept=concept,
                    rank=i,
                    show_citations=config.enable_citations,
                    show_related=True
                )
        
        # Tab 2: Relaciones entre conceptos
        with tab2:
            st.markdown("### 🔗 Relaciones entre Conceptos")
            
            st.markdown("""
            <div style="background: #34495e; padding: 1rem; border-radius: 8px; border-left: 4px solid #3498db; color: #ecf0f1; margin: 1rem 0;">
                <strong>ℹ️ Sobre las relaciones:</strong> Los conceptos relacionados son aquellos que
                aparecen frecuentemente juntos en los mismos fragmentos de texto. Esto puede indicar
                conexiones temáticas importantes.
            </div>
            """, unsafe_allow_html=True)
            
            # Seleccionar concepto para ver relaciones
            concept_names = [c.concept for c in concepts[:30]]  # Top 30
            selected_concept_name = st.selectbox(
                "Selecciona un concepto para ver sus relaciones",
                concept_names,
                help="Muestra conceptos que co-ocurren con el seleccionado",
                key="concepts_relation_selector"
            )
            
            # Encontrar concepto seleccionado
            selected_concept = next(c for c in concepts if c.concept == selected_concept_name)
            
            # Mostrar información del concepto
            col1, col2 = st.columns([2, 1])
            
            with col1:
                st.markdown(f"### 📌 {selected_concept.concept}")
                st.markdown(f"""
                - **Frecuencia:** {selected_concept.frequency}
                - **Relevancia:** {selected_concept.relevance_score:.3f}
                - **Fuentes:** {len(selected_concept.sources)}
                """)
            
            with col2:
                st.metric(
                    "🔗 Conceptos Relacionados",
                    len(selected_concept.related_concepts)
                )
            
            # Mostrar conceptos relacionados
            if selected_concept.related_concepts:
                st.markdown("#### Conceptos que aparecen junto a este:")
                
                for i, related_name in enumerate(selected_concept.related_concepts, 1):
                    # Encontrar información del concepto relacionado
                    related_concept = next((c for c in concepts if c.concept == related_name), None)
                    
                    if related_concept:
                        st.markdown(f"""
                        <div style="background: #2c3e50; padding: 0.8rem; border-radius: 5px; margin: 0.5rem 0; border-left: 3px solid #3498db; color: #ecf0f1;">
                            <strong>{i}. {related_concept.concept}</strong><br>
                            <small>Frecuencia: {related_concept.frequency} | Relevancia: {related_concept.relevance_score:.3f}</small>
                        </div>
                        """, unsafe_allow_html=True)
            else:
                st.info("Este concepto no tiene relaciones fuertes identificadas")
            
            # Matriz de co-ocurrencia (simplificada)
            st.markdown("---")
            st.markdown("#### 📊 Matriz de Co-ocurrencia (Top 10)")
            
            # Crear matriz simple
            top_10 = concepts[:10]
            matrix_data = []
            
            for concept in top_10:
                row = {
                    'Concepto': concept.concept,
                    'Relaciones': len(concept.related_concepts)
                }
                matrix_data.append(row)
            
            st.dataframe(matrix_data, use_container_width=True)
        
        # Tab 3: Bibliografía
        with tab3:
            st.markdown("### 📚 Bibliografía y Referencias")
            
            st.markdown("""
            <div style="background: #34495e; padding: 1rem; border-radius: 8px; border-left: 4px solid #9b59b6; color: #ecf0f1; margin: 1rem 0;">
                <strong>ℹ️ Fundamentación científica:</strong> Cada concepto está respaldado por citas
                específicas a las fuentes originales. Esta sección te permite ver todas las referencias
                y verificar la fundamentación del análisis.
            </div>
            """, unsafe_allow_html=True)
            
            # Estadísticas de citación
            cit_stats = extractor.citation_manager.get_statistics()
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("📖 Total de Citas", cit_stats['total_citations'])
            with col2:
                st.metric("📂 Fuentes Citadas", cit_stats['unique_sources'])
            with col3:
                st.metric("⭐ Relevancia Promedio", f"{cit_stats['avg_relevance']:.3f}")
            
            # Bibliografía
            st.markdown("#### 📋 Referencias Bibliográficas")
            
            bibliography = extractor.citation_manager.generate_bibliography()
            
            for i, ref in enumerate(bibliography, 1):
                st.markdown(f"{i}. {ref}")
            
            # Distribución de citas por fuente
            if 'citations_by_source' in cit_stats:
                st.markdown("---")
                st.markdown("#### 📊 Distribución de Citas por Fuente")
                
                import plotly.graph_objects as go
                
                sources_list = list(cit_stats['citations_by_source'].keys())
                counts = list(cit_stats['citations_by_source'].values())
                
                # Truncar nombres largos
                sources_display = [s.split('/')[-1][:30] for s in sources_list]
                
                fig = go.Figure([go.Bar(
                    x=counts,
                    y=sources_display,
                    orientation='h',
                    marker_color='#00FF99'
                )])
                
                fig.update_layout(
                    title="Número de Citas por Fuente",
                    xaxis_title="Número de Citas",
                    yaxis_title="Fuente",
                    height=max(400, len(sources_list) * 30),
                    template="plotly_dark"
                )
                
                st.plotly_chart(fig, use_container_width=True)
        
        # Tab 4: Exportar
        with tab4:
            st.markdown("### 💾 Exportar a Word")
            
            st.markdown("""
            <div style="background: #34495e; padding: 1rem; border-radius: 8px; border-left: 4px solid #e67e22; color: #ecf0f1; margin: 1rem 0;">
                <strong>ℹ️ Exportación a Word:</strong> Genera un documento de Word profesional con todos los conceptos extraídos,
                incluyendo explicaciones, categorías y citas para uso en tu investigación.
            </div>
            """, unsafe_allow_html=True)
            
            # Opciones de exportación
            col1, col2 = st.columns(2)
            
            with col1:
                include_explanations = st.checkbox(
                    "Incluir explicaciones de conceptos",
                    value=True,
                    help="Incluye las explicaciones generadas por la IA para cada concepto",
                    key="concepts_include_explanations"
                )
            
            with col2:
                include_citations = st.checkbox(
                    "Incluir citas y referencias",
                    value=True,
                    help="Incluye las citas específicas donde aparece cada concepto",
                    key="concepts_include_citations"
                )
            
            # Información del documento
            st.markdown("#### 📄 Información del Documento")
            
            col1, col2 = st.columns(2)
            with col1:
                document_title = st.text_input(
                    "Título del documento",
                    value="Análisis de Conceptos Clave",
                    help="Título que aparecerá en el documento",
                    key="concepts_document_title"
                )
            
            with col2:
                author_name = st.text_input(
                    "Autor",
                    value="Investigador",
                    help="Nombre del autor del análisis",
                    key="concepts_author_name"
                )
            
            # Generar documento Word
            if st.button("📄 Generar Documento Word", type="primary", use_container_width=True, key="concepts_generate_word"):
                try:
                    # Crear documento Word
                    doc_content = _generate_word_document(
                        concepts=concepts,
                        summary=summary,
                        title=document_title,
                        author=author_name,
                        include_explanations=include_explanations,
                        include_citations=include_citations
                    )
                    
                    st.download_button(
                        label="💾 Descargar Documento Word",
                        data=doc_content,
                        file_name=f"{document_title.replace(' ', '_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                    st.success("✅ Documento Word generado correctamente")
                    
                except Exception as e:
                    st.error(f"❌ Error al generar documento Word: {str(e)}")
                    st.exception(e)


def _generate_word_document(
    concepts: List[ExtractedConcept],
    summary: Dict[str, Any],
    title: str,
    author: str,
    include_explanations: bool = True,
    include_citations: bool = True
) -> bytes:
    """
    Generar documento Word con los conceptos extraídos
    
    Args:
        concepts: Lista de conceptos extraídos
        summary: Resumen estadístico
        title: Título del documento
        author: Autor del documento
        include_explanations: Si incluir explicaciones
        include_citations: Si incluir citas
        
    Returns:
        Contenido del documento Word como bytes
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        from docx.oxml.shared import OxmlElement, qn
        import io
    except ImportError:
        raise ImportError("python-docx es requerido para generar documentos Word. Instala con: pip install python-docx")
    
    # Crear documento
    doc = Document()
    
    # Configurar estilos
    title_style = doc.styles['Title']
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    
    heading_style = doc.styles['Heading 1']
    heading_style.font.size = Pt(14)
    heading_style.font.bold = True
    
    # Título del documento
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Información del documento
    doc.add_paragraph(f"Autor: {author}")
    doc.add_paragraph(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph("")
    
    # Resumen ejecutivo
    doc.add_heading("Resumen Ejecutivo", level=1)
    
    summary_text = f"""
    Este documento presenta un análisis cualitativo de conceptos clave extraídos de documentos de investigación.
    
    Estadísticas del análisis:
    • Total de conceptos identificados: {summary.get('total_concepts', 0)}
    • Fuentes analizadas: {summary.get('unique_sources', 0)}
    • Total de citas generadas: {summary.get('total_citations', 0)}
    • Relevancia promedio: {summary.get('avg_relevance', 0):.3f}
    """
    
    doc.add_paragraph(summary_text)
    doc.add_paragraph("")
    
    # Conceptos principales
    doc.add_heading("Conceptos Clave Identificados", level=1)
    
    for i, concept in enumerate(concepts, 1):
        # Título del concepto
        concept_title = f"{i}. {concept.concept}"
        doc.add_heading(concept_title, level=2)
        
        # Información básica
        info_text = f"""
        Frecuencia: {concept.frequency} apariciones
        Puntuación de relevancia: {concept.relevance_score:.3f}
        Fuentes: {len(concept.sources)}
        """
        
        if hasattr(concept, 'category') and concept.category:
            info_text += f"Categoría: {concept.category}\n"
        
        doc.add_paragraph(info_text)
        
        # Explicación del concepto
        if include_explanations and concept.context_examples:
            explanation = None
            for example in concept.context_examples:
                if "Explicación:" in example:
                    explanation = example.replace("Explicación:", "").strip()
                    break
            
            if explanation:
                doc.add_heading("Explicación", level=3)
                doc.add_paragraph(explanation)
        
        # Citas y referencias
        if include_citations and concept.citations:
            doc.add_heading("Referencias", level=3)
            
            for j, citation in enumerate(concept.citations[:3], 1):  # Máximo 3 citas
                citation_text = f"""
                Cita {j}:
                Fuente: {citation.source_file}
                Contexto: {citation.get_full_context(max_chars=200)}
                """
                doc.add_paragraph(citation_text)
        
        # Conceptos relacionados
        if concept.related_concepts:
            doc.add_heading("Conceptos Relacionados", level=3)
            related_text = ", ".join(concept.related_concepts)
            doc.add_paragraph(related_text)
        
        doc.add_paragraph("")  # Espacio entre conceptos
    
    # Bibliografía
    if include_citations:
        doc.add_heading("Bibliografía", level=1)
        
        # Obtener bibliografía del extractor
        from ...extractors.concept_extractor import ConceptExtractor
        extractor = ConceptExtractor()
        bibliography = extractor.citation_manager.generate_bibliography()
        
        for i, ref in enumerate(bibliography, 1):
            doc.add_paragraph(f"{i}. {ref}")
    
    # Guardar en bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes.getvalue()

