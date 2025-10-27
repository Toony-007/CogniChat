"""
Extractor de Temas para Análisis Cualitativo
Implementa LDA, BERTopic y clustering semántico para identificación de temas
"""

import re
import numpy as np
from typing import List, Dict, Any, Tuple, Optional
from dataclasses import dataclass
from datetime import datetime
import logging

# Importaciones para procesamiento de texto
from sklearn.feature_extraction.text import TfidfVectorizer, CountVectorizer
from sklearn.decomposition import LatentDirichletAllocation
from sklearn.cluster import KMeans
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.manifold import TSNE
from sklearn.decomposition import PCA

# Importaciones para NLP
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from nltk.stem import SnowballStemmer

# Importaciones para visualización
import plotly.express as px
import plotly.graph_objects as go
from wordcloud import WordCloud

# Importaciones del sistema
from ..core.config import AnalysisConfig
from ..core.citation_manager import CitationManager
from utils.ollama_client import ollama_client

logger = logging.getLogger(__name__)

@dataclass
class TopicData:
    """Datos de un tema identificado"""
    topic_id: int
    topic_name: str
    keywords: List[str]
    coherence_score: float
    frequency: int
    documents: List[str]
    citations: List[Dict[str, Any]]
    description: str
    confidence: float

@dataclass
class TopicAnalysisResult:
    """Resultado del análisis de temas"""
    topics: List[TopicData]
    algorithm_used: str
    total_topics: int
    coherence_avg: float
    processing_time: float
    metadata: Dict[str, Any]
    citations: List[Dict[str, Any]]

class TopicExtractor:
    """
    Extractor de temas especializado para análisis cualitativo
    
    Características:
    - Múltiples algoritmos (LDA, BERTopic, Clustering)
    - Sistema de citación integrado
    - Análisis de coherencia
    - Generación de descripciones contextuales
    """
    
    def __init__(self, config: AnalysisConfig):
        self.config = config
        self.citation_manager = CitationManager()
        self.stemmer = SnowballStemmer('spanish')
        self.stopwords = self._get_spanish_stopwords()
        
        # Inicializar NLTK si es necesario
        self._initialize_nltk()
    
    def extract_topics_hybrid(self, chunks: List[Dict[str, Any]]) -> TopicAnalysisResult:
        """
        Extraer temas usando enfoque híbrido: Algoritmo + IA
        
        Proceso híbrido:
        1. Extracción inicial con algoritmo (LDA o Clustering)
        2. Refinamiento con IA para mejorar nombres y explicaciones
        3. Validación de calidad de temas
        4. Enriquecimiento con citaciones
        
        Args:
            chunks: Lista de chunks de texto
            
        Returns:
            Resultado del análisis de temas híbrido
        """
        print("🎯 Iniciando análisis híbrido de temas...")
        
        # FASE 1: Extracción inicial con algoritmo
        print(f"🔍 Fase 1: Extracción inicial con {self.config.topic_algorithm.upper()}...")
        
        if self.config.topic_algorithm == "lda":
            raw_result = self.extract_topics_lda(chunks, self.config.max_topics)
        elif self.config.topic_algorithm == "clustering":
            raw_result = self.extract_topics_clustering(chunks, self.config.max_topics)
        else:
            raise ValueError(f"Algoritmo desconocido: {self.config.topic_algorithm}")
        
        # FASE 2: Refinamiento con IA (si está habilitado)
        if self.config.enable_topic_refinement:
            print(f"🤖 Fase 2: Refinamiento con {self.config.llm_model}...")
            refined_result = self._refine_topics_with_llm(raw_result, chunks)
        else:
            print("⚠️ Refinamiento con IA deshabilitado, usando resultados algorítmicos")
            refined_result = raw_result
        
        print(f"✅ Análisis híbrido completado: {len(refined_result.topics)} temas identificados")
        return refined_result
    
    def _refine_topics_with_llm(self, raw_result: TopicAnalysisResult, chunks: List[Dict[str, Any]]) -> TopicAnalysisResult:
        """
        Refinar temas usando modelo LLM (DeepSeek R1)
        
        Este método toma los temas candidatos extraídos por algoritmos
        y los envía al LLM para:
        1. Mejorar nombres de temas
        2. Generar explicaciones detalladas
        3. Filtrar temas irrelevantes
        4. Categorizar temas
        
        Args:
            raw_result: Resultado del análisis algorítmico
            chunks: Chunks originales para contexto
            
        Returns:
            Resultado refinado con IA
        """
        try:
            # Preparar datos para el LLM
            topics_data = []
            for topic in raw_result.topics:
                topics_data.append({
                    'id': topic.topic_id,
                    'name': topic.topic_name,
                    'keywords': topic.keywords[:10],  # Top 10 palabras clave
                    'frequency': topic.frequency,
                    'coherence': topic.coherence_score,
                    'description': topic.description
                })
            
            # Crear contexto del documento
            document_context = " ".join([chunk.get('content', '') for chunk in chunks[:10]])  # Primeros 10 chunks
            document_context = document_context[:3000] + "..." if len(document_context) > 3000 else document_context
            
            # Crear prompt para el LLM
            prompt = self._create_topic_refinement_prompt(topics_data, document_context)
            
            # Llamar al LLM sin límites de tokens
            response = ollama_client.generate_response(
                model=self.config.llm_model,
                prompt=prompt
            )
            
            # Parsear respuesta del LLM
            refined_topics = self._parse_llm_topic_response(response, raw_result.topics)
            
            # Crear resultado refinado
            refined_result = TopicAnalysisResult(
                topics=refined_topics,
                algorithm_used=f"{raw_result.algorithm_used} + {self.config.llm_model}",
                total_topics=len(refined_topics),
                coherence_avg=raw_result.coherence_avg,
                processing_time=raw_result.processing_time,
                metadata={
                    **raw_result.metadata,
                    'llm_refinement': True,
                    'llm_model': self.config.llm_model
                },
                citations=raw_result.citations
            )
            
            return refined_result
            
        except Exception as e:
            print(f"Error en refinamiento LLM: {e}")
            # Fallback: devolver resultado original
            return raw_result
    
    def _create_topic_refinement_prompt(self, topics_data: List[Dict], document_context: str) -> str:
        """
        Crear prompt optimizado para refinamiento de temas con DeepSeek R1
        
        Args:
            topics_data: Datos de temas candidatos
            document_context: Contexto del documento
            
        Returns:
            Prompt formateado para el LLM
        """
        topics_section = ""
        for topic in topics_data:
            topics_section += f"""
Tema {topic['id'] + 1}:
- Nombre actual: {topic['name']}
- Palabras clave: {', '.join(topic['keywords'][:8])}
- Frecuencia: {topic['frequency']}
- Coherencia: {topic['coherence']:.3f}
- Descripción: {topic['description']}
"""
        
        return f"""Eres un experto en análisis cualitativo de documentos académicos. Tu tarea es refinar y mejorar los temas identificados por algoritmos estadísticos.

IMPORTANTE: Responde ÚNICAMENTE en español. Todos los temas, nombres y explicaciones deben estar en español.

CONTEXTO DEL DOCUMENTO:
{document_context}

TEMAS CANDIDATOS IDENTIFICADOS POR ALGORITMO:
{topics_section}

INSTRUCCIONES CRÍTICAS:
1. ANALIZA cada tema candidato y determina si es relevante y significativo
2. MEJORA los nombres de temas para que sean más descriptivos y académicos
3. GENERA explicaciones detalladas de qué representa cada tema
4. FILTRA temas que sean irrelevantes, genéricos o de baja calidad
5. CATEGORIZA cada tema (metodología, teoría, resultado, proceso, fenómeno social)
6. PRIORIZA temas que sean útiles para investigación cualitativa
7. CADA explicación debe ser ÚNICA y específica al tema

FORMATO DE RESPUESTA (JSON estricto):
{{
    "temas_refinados": [
        {{
            "id": 0,
            "nombre": "nombre_del_tema_académico_profundo_en_español",
            "categoria": "metodologia|teoria|resultado|proceso|fenomeno_social",
            "explicacion": "Explicación detallada y ÚNICA de qué representa este tema específico y por qué es importante para la investigación.",
            "palabras_clave": ["palabra1", "palabra2", "palabra3"],
            "relevancia": "alta|media|baja",
            "justificacion": "Por qué este tema es relevante y qué fenómeno captura"
        }}
    ]
}}

CRITERIOS DE CALIDAD:
- Los temas deben ser académicos y significativos
- Deben capturar fenómenos, procesos o conceptos importantes
- Deben ser específicos al contexto del documento
- NO deben ser temas genéricos o irrelevantes
- CADA explicación debe ser ÚNICA y específica al tema
- Los nombres deben ser descriptivos y profesionales
- Máximo {self.config.max_topics} temas de alta calidad

EJEMPLOS DE TEMAS VÁLIDOS:
✅ "Metodología de investigación cualitativa" (específico, académico)
✅ "Procesos de aprendizaje colaborativo" (complejo, significativo)
✅ "Factores de resiliencia comunitaria" (específico, social)

EJEMPLOS DE TEMAS INVÁLIDOS:
❌ "Texto general" (genérico)
❌ "Palabras comunes" (irrelevante)
❌ "Documento" (genérico)

REGLAS CRÍTICAS:
- SOLO incluye temas que sean académicamente relevantes
- Cada explicación debe ser completamente diferente
- Las explicaciones deben ser específicas al tema, no genéricas
- NO uses plantillas o frases genéricas repetidas
- Si no encuentras temas relevantes, devuelve menos temas pero de mayor calidad

IMPORTANTE:
- Responde SOLO con JSON válido
- Máximo {self.config.max_topics} temas
- Cada tema debe ser una idea compleja y significativa
- TODO debe estar en español"""
    
    def _parse_llm_topic_response(self, response: str, original_topics: List[TopicData]) -> List[TopicData]:
        """
        Parsear respuesta del LLM y crear temas refinados
        
        Args:
            response: Respuesta del LLM
            original_topics: Temas originales para preservar metadatos
            
        Returns:
            Lista de temas refinados
        """
        import json
        import re
        
        try:
            # Limpiar caracteres de control inválidos
            cleaned_response = self._clean_json_response(response)
            
            # Buscar JSON en la respuesta
            json_start = cleaned_response.find('{')
            json_end = cleaned_response.rfind('}') + 1
            
            if json_start == -1 or json_end == 0:
                raise ValueError("No se encontró JSON válido en la respuesta")
            
            json_str = cleaned_response[json_start:json_end]
            
            # Verificar si el JSON está completo
            if not self._is_json_complete(json_str):
                print("⚠️ JSON parece estar truncado, intentando reparar...")
                json_str = self._repair_truncated_json(json_str)
            
            # Intentar parsear JSON
            try:
                data = json.loads(json_str)
            except json.JSONDecodeError as e:
                print(f"⚠️ Error JSON inicial: {e}")
                json_str = self._aggressive_json_clean(json_str)
                try:
                    data = json.loads(json_str)
                    print("✅ JSON parseado exitosamente después de limpieza")
                except json.JSONDecodeError as e2:
                    print(f"❌ Error JSON persistente: {e2}")
                    raise ValueError(f"No se pudo parsear JSON: {e2}")
            
            refined_topics = []
            
            for item in data.get('temas_refinados', []):
                # Buscar tema original correspondiente
                original_topic = None
                for orig_topic in original_topics:
                    if orig_topic.topic_id == item.get('id', -1):
                        original_topic = orig_topic
                        break
                
                if not original_topic:
                    continue
                
                # Crear tema refinado
                refined_topic = TopicData(
                    topic_id=item.get('id', original_topic.topic_id),
                    topic_name=item.get('nombre', original_topic.topic_name),
                    keywords=item.get('palabras_clave', original_topic.keywords),
                    coherence_score=original_topic.coherence_score,
                    frequency=original_topic.frequency,
                    documents=original_topic.documents,
                    citations=original_topic.citations,
                    description=item.get('explicacion', original_topic.description),
                    confidence=original_topic.confidence
                )
                
                # Agregar información adicional del LLM
                if self.config.include_topic_explanations and 'explicacion' in item:
                    refined_topic.description = item['explicacion']
                    print(f"✅ Explicación refinada para '{item['nombre']}': {item['explicacion'][:100]}...")
                
                refined_topics.append(refined_topic)
            
            return refined_topics
            
        except (json.JSONDecodeError, KeyError, ValueError) as e:
            print(f"Error parseando respuesta LLM: {e}")
            print(f"Respuesta original: {response[:500]}...")
            # Fallback: devolver temas originales
            return original_topics
    
    def _clean_json_response(self, response: str) -> str:
        """Limpiar caracteres de control inválidos de la respuesta del LLM"""
        import re
        cleaned = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', response)
        cleaned = cleaned.replace('\r\n', '\n').replace('\r', '\n')
        cleaned = re.sub(r' +', ' ', cleaned)
        return cleaned
    
    def _is_json_complete(self, json_str: str) -> bool:
        """Verificar si el JSON está completo (no truncado)"""
        open_braces = json_str.count('{')
        close_braces = json_str.count('}')
        open_brackets = json_str.count('[')
        close_brackets = json_str.count(']')
        
        braces_balanced = open_braces == close_braces
        brackets_balanced = open_brackets == close_brackets
        ends_properly = json_str.strip().endswith('}') or json_str.strip().endswith(']')
        
        return braces_balanced and brackets_balanced and ends_properly
    
    def _repair_truncated_json(self, json_str: str) -> str:
        """Reparar JSON truncado agregando elementos faltantes"""
        import re
        
        if not json_str.strip().endswith('}'):
            last_complete_brace = json_str.rfind('},')
            if last_complete_brace != -1:
                json_str = json_str[:last_complete_brace + 1] + '\n    ]\n}'
            else:
                json_str = json_str.rstrip(',') + '\n    ]\n}'
        
        json_str = re.sub(r',\s*,', ',', json_str)
        json_str = re.sub(r',\s*}', '}', json_str)
        json_str = re.sub(r',\s*]', ']', json_str)
        
        return json_str
    
    def _aggressive_json_clean(self, json_str: str) -> str:
        """Limpieza agresiva de JSON para manejar casos problemáticos"""
        import re
        
        def escape_string(match):
            string_content = match.group(1)
            string_content = string_content.replace('\\', '\\\\')
            string_content = string_content.replace('"', '\\"')
            string_content = string_content.replace('\n', '\\n')
            string_content = string_content.replace('\t', '\\t')
            string_content = string_content.replace('\r', '\\r')
            return f'"{string_content}"'
        
        json_str = re.sub(r'"([^"]*)"', escape_string, json_str)
        json_str = re.sub(r'"\s*}\s*', '",\n}', json_str)
        json_str = re.sub(r'"\s*]\s*', '",\n]', json_str)
        json_str = re.sub(r'"\s*"', '",\n"', json_str)
        json_str = re.sub(r',\s*,', ',', json_str)
        json_str = re.sub(r',\s*}', '}', json_str)
        json_str = re.sub(r',\s*]', ']', json_str)
        json_str = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', json_str)
        
        return json_str
    
    def _initialize_nltk(self):
        """Inicializar recursos de NLTK"""
        try:
            nltk.data.find('tokenizers/punkt')
        except LookupError:
            nltk.download('punkt', quiet=True)
        
        try:
            nltk.data.find('corpora/stopwords')
        except LookupError:
            nltk.download('stopwords', quiet=True)
    
    def _get_spanish_stopwords(self) -> List[str]:
        """Obtener stopwords en español"""
        try:
            spanish_stopwords = set(stopwords.words('spanish'))
            # Agregar stopwords personalizados
            custom_stopwords = {
                'también', 'puede', 'ser', 'está', 'están', 'hacer', 'hace',
                'tiene', 'tienen', 'dice', 'dice', 'según', 'mismo', 'misma',
                'mismos', 'mismas', 'otro', 'otra', 'otros', 'otras', 'cada',
                'todo', 'toda', 'todos', 'todas', 'muy', 'más', 'menos',
                'mucho', 'muchos', 'poco', 'pocos', 'algunos', 'algunas'
            }
            return list(spanish_stopwords.union(custom_stopwords))
        except:
            return []
    
    def extract_topics_lda(self, chunks: List[Dict[str, Any]], n_topics: int = 10) -> TopicAnalysisResult:
        """
        Extraer temas usando LDA (Latent Dirichlet Allocation)
        
        Args:
            chunks: Lista de chunks de texto
            n_topics: Número de temas a identificar
            
        Returns:
            Resultado del análisis de temas
        """
        start_time = datetime.now()
        
        try:
            # Preprocesar textos
            texts = self._preprocess_texts(chunks)
            
            if len(texts) < n_topics:
                raise ValueError(f"Insuficientes documentos para {n_topics} temas. Mínimo requerido: {n_topics}")
            
            # Vectorización TF-IDF
            vectorizer = TfidfVectorizer(
                max_features=1000,
                stop_words=self.stopwords,
                ngram_range=(1, 2),
                min_df=2,
                max_df=0.8
            )
            
            tfidf_matrix = vectorizer.fit_transform(texts)
            feature_names = vectorizer.get_feature_names_out()
            
            # Verificar que la matriz no esté vacía
            if tfidf_matrix.shape[0] == 0 or tfidf_matrix.shape[1] == 0:
                raise ValueError("La matriz TF-IDF está vacía. Verifica que hay suficiente contenido de texto.")
            
            # Aplicar LDA
            lda = LatentDirichletAllocation(
                n_components=n_topics,
                random_state=42,
                max_iter=100,
                learning_decay=0.7
            )
            
            lda.fit(tfidf_matrix)
            
            # Extraer temas
            topics = self._extract_lda_topics(lda, feature_names, n_topics)
            
            # Calcular coherencia
            coherence_scores = self._calculate_topic_coherence(topics, texts)
            
            # Enriquecer con citaciones
            enriched_topics = self._enrich_topics_with_citations(topics, chunks, coherence_scores)
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return TopicAnalysisResult(
                topics=enriched_topics,
                algorithm_used="LDA",
                total_topics=len(enriched_topics),
                coherence_avg=np.mean(coherence_scores),
                processing_time=processing_time,
                metadata={
                    'n_topics': n_topics,
                    'n_documents': len(chunks),
                    'n_features': len(feature_names)
                },
                citations=self.citation_manager.export_citations()
            )
            
        except Exception as e:
            logger.error(f"Error en extracción LDA: {str(e)}")
            raise
    
    def extract_topics_clustering(self, chunks: List[Dict[str, Any]], n_clusters: int = 10) -> TopicAnalysisResult:
        """
        Extraer temas usando clustering semántico
        
        Args:
            chunks: Lista de chunks de texto
            n_clusters: Número de clusters/temas
            
        Returns:
            Resultado del análisis de temas
        """
        start_time = datetime.now()
        
        try:
            # Preprocesar textos
            texts = self._preprocess_texts(chunks)
            
            if len(texts) < n_clusters:
                raise ValueError(f"Insuficientes documentos para {n_clusters} clusters. Mínimo requerido: {n_clusters}")
            
            # Vectorización TF-IDF
            vectorizer = TfidfVectorizer(
                max_features=1000,
                stop_words=self.stopwords,
                ngram_range=(1, 2),
                min_df=2,
                max_df=0.8
            )
            
            tfidf_matrix = vectorizer.fit_transform(texts)
            feature_names = vectorizer.get_feature_names_out()
            
            # Verificar que la matriz no esté vacía
            if tfidf_matrix.shape[0] == 0 or tfidf_matrix.shape[1] == 0:
                raise ValueError("La matriz TF-IDF está vacía. Verifica que hay suficiente contenido de texto.")
            
            # Aplicar K-Means clustering
            kmeans = KMeans(n_clusters=n_clusters, random_state=42, n_init=10)
            cluster_labels = kmeans.fit_predict(tfidf_matrix)
            
            # Extraer temas de clusters
            topics = self._extract_cluster_topics(
                kmeans, tfidf_matrix, feature_names, cluster_labels, texts
            )
            
            # Calcular coherencia
            coherence_scores = self._calculate_topic_coherence(topics, texts)
            
            # Enriquecer con citaciones
            enriched_topics = self._enrich_topics_with_citations(topics, chunks, coherence_scores)
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return TopicAnalysisResult(
                topics=enriched_topics,
                algorithm_used="Clustering Semántico",
                total_topics=len(enriched_topics),
                coherence_avg=np.mean(coherence_scores),
                processing_time=processing_time,
                metadata={
                    'n_clusters': n_clusters,
                    'n_documents': len(chunks),
                    'n_features': len(feature_names)
                },
                citations=self.citation_manager.export_citations()
            )
            
        except Exception as e:
            logger.error(f"Error en clustering semántico: {str(e)}")
            raise
    
    def _preprocess_texts(self, chunks: List[Dict[str, Any]]) -> List[str]:
        """Preprocesar textos para análisis de temas"""
        texts = []
        
        for chunk in chunks:
            if isinstance(chunk, dict) and 'content' in chunk:
                text = chunk['content']
                
                # Limpiar texto
                text = re.sub(r'[^\w\s]', ' ', text)
                text = re.sub(r'\s+', ' ', text)
                text = text.strip().lower()
                
                if len(text) > 50:  # Solo textos significativos
                    texts.append(text)
        
        return texts
    
    def _extract_lda_topics(self, lda_model, feature_names: List[str], n_topics: int) -> List[TopicData]:
        """Extraer temas del modelo LDA"""
        topics = []
        
        for topic_idx, topic in enumerate(lda_model.components_):
            # Obtener palabras más importantes
            top_words_idx = topic.argsort()[-10:][::-1]
            top_words = [feature_names[i] for i in top_words_idx]
            top_scores = [topic[i] for i in top_words_idx]
            
            # Crear nombre del tema
            topic_name = f"Tema {topic_idx + 1}: {' '.join(top_words[:3])}"
            
            # Calcular frecuencia
            frequency = int(np.sum(topic))
            
            # Crear descripción
            description = self._generate_topic_description(top_words, top_scores)
            
            topics.append(TopicData(
                topic_id=topic_idx,
                topic_name=topic_name,
                keywords=top_words,
                coherence_score=0.0,  # Se calculará después
                frequency=frequency,
                documents=[],  # Se llenará después
                citations=[],  # Se llenará después
                description=description,
                confidence=float(np.mean(top_scores))
            ))
        
        return topics
    
    def _extract_cluster_topics(self, kmeans_model, tfidf_matrix, feature_names: List[str], 
                               cluster_labels: np.ndarray, texts: List[str]) -> List[TopicData]:
        """Extraer temas de clusters"""
        topics = []
        
        for cluster_id in range(kmeans_model.n_clusters):
            # Obtener documentos del cluster
            cluster_mask = cluster_labels == cluster_id
            cluster_docs = tfidf_matrix[cluster_mask]
            
            if cluster_docs.shape[0] == 0:
                continue
            
            # Calcular centroide del cluster
            centroid = np.mean(cluster_docs, axis=0).A1
            
            # Obtener palabras más importantes
            top_words_idx = centroid.argsort()[-10:][::-1]
            top_words = [feature_names[i] for i in top_words_idx]
            top_scores = [centroid[i] for i in top_words_idx]
            
            # Crear nombre del tema
            topic_name = f"Cluster {cluster_id + 1}: {' '.join(top_words[:3])}"
            
            # Calcular frecuencia
            frequency = int(np.sum(cluster_mask))
            
            # Crear descripción
            description = self._generate_topic_description(top_words, top_scores)
            
            topics.append(TopicData(
                topic_id=cluster_id,
                topic_name=topic_name,
                keywords=top_words,
                coherence_score=0.0,  # Se calculará después
                frequency=frequency,
                documents=[],  # Se llenará después
                citations=[],  # Se llenará después
                description=description,
                confidence=float(np.mean(top_scores))
            ))
        
        return topics
    
    def _calculate_topic_coherence(self, topics: List[TopicData], texts: List[str]) -> List[float]:
        """Calcular coherencia de temas"""
        coherence_scores = []
        
        for topic in topics:
            # Calcular coherencia basada en co-ocurrencia de palabras
            coherence = self._calculate_topic_coherence_score(topic.keywords, texts)
            coherence_scores.append(coherence)
        
        return coherence_scores
    
    def _calculate_topic_coherence_score(self, keywords: List[str], texts: List[str]) -> float:
        """Calcular score de coherencia mejorado para un tema"""
        try:
            if len(keywords) < 2:
                return 0.0
            
            # Usar las primeras 10 palabras clave para mejor cálculo
            top_keywords = keywords[:10]
            
            # Calcular coherencia usando ventana deslizante
            coherence_scores = []
            
            for text in texts:
                words = text.split()
                if len(words) < 2:
                    continue
                
                # Ventana deslizante de tamaño 10
                window_size = min(10, len(words))
                
                for i in range(len(words) - window_size + 1):
                    window = words[i:i + window_size]
                    
                    # Contar co-ocurrencias en la ventana
                    window_keywords = [w for w in window if w in top_keywords]
                    
                    if len(window_keywords) >= 2:
                        # Calcular coherencia local
                        unique_keywords = set(window_keywords)
                        if len(unique_keywords) > 1:
                            coherence_scores.append(len(unique_keywords) / len(top_keywords))
            
            if coherence_scores:
                # Promedio de coherencias locales
                avg_coherence = sum(coherence_scores) / len(coherence_scores)
                
                # Normalizar entre 0 y 1
                normalized_coherence = min(1.0, avg_coherence * 2)  # Factor de escala
                
                return normalized_coherence
            else:
                return 0.0
                
        except Exception as e:
            print(f"⚠️ Error calculando coherencia: {e}")
            return 0.0
    
    def _generate_topic_description(self, keywords: List[str], scores: List[float]) -> str:
        """Generar descripción contextual del tema"""
        try:
            # Crear descripción basada en palabras clave
            top_keywords = keywords[:5]
            description = f"Tema caracterizado por: {', '.join(top_keywords)}"
            
            # Agregar contexto basado en palabras
            if any(word in ['educación', 'formación', 'aprendizaje'] for word in top_keywords):
                description += ". Este tema se relaciona con aspectos educativos y formativos."
            elif any(word in ['trabajo', 'empleo', 'laboral'] for word in top_keywords):
                description += ". Este tema se relaciona con aspectos laborales y profesionales."
            elif any(word in ['social', 'comunidad', 'sociedad'] for word in top_keywords):
                description += ". Este tema se relaciona con aspectos sociales y comunitarios."
            
            return description
            
        except Exception:
            return f"Tema caracterizado por: {', '.join(keywords[:3])}"
    
    def _enrich_topics_with_citations(self, topics: List[TopicData], chunks: List[Dict[str, Any]], 
                                    coherence_scores: List[float]) -> List[TopicData]:
        """Enriquecer temas con citaciones y documentos"""
        enriched_topics = []
        
        for i, topic in enumerate(topics):
            # Actualizar coherencia
            topic.coherence_score = coherence_scores[i] if i < len(coherence_scores) else 0.0
            
            # Buscar documentos relacionados
            related_docs = self._find_related_documents(topic, chunks)
            topic.documents = related_docs
            
            # Generar citaciones
            citations = self._generate_topic_citations(topic, chunks)
            topic.citations = citations
            
            enriched_topics.append(topic)
        
        return enriched_topics
    
    def _find_related_documents(self, topic: TopicData, chunks: List[Dict[str, Any]]) -> List[str]:
        """Encontrar documentos relacionados con el tema"""
        related_docs = []
        
        for chunk in chunks:
            if isinstance(chunk, dict) and 'content' in chunk:
                content = chunk['content'].lower()
                
                # Verificar si el chunk contiene palabras clave del tema
                keyword_matches = sum(1 for keyword in topic.keywords[:5] 
                                    if keyword.lower() in content)
                
                if keyword_matches >= 2:  # Al menos 2 palabras clave
                    source = chunk.get('metadata', {}).get('source_file', 'unknown')
                    if source not in related_docs:
                        related_docs.append(source)
        
        return related_docs
    
    def _generate_topic_citations(self, topic: TopicData, chunks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Generar citaciones para el tema"""
        citations = []
        
        for chunk in chunks:
            if isinstance(chunk, dict) and 'content' in chunk:
                content = chunk['content']
                
                # Verificar si el chunk contiene palabras clave del tema
                keyword_matches = sum(1 for keyword in topic.keywords[:5] 
                                    if keyword.lower() in content.lower())
                
                if keyword_matches >= 2:
                    citation = self.citation_manager.add_citation(
                        source_file=chunk.get('metadata', {}).get('source_file', 'unknown'),
                        chunk_id=chunk.get('metadata', {}).get('chunk_index', 0),
                        content=content,
                        context_before="",  # Se puede mejorar después
                        context_after="",  # Se puede mejorar después
                        relevance_score=keyword_matches / len(topic.keywords[:5])
                    )
                    citations.append(citation.to_dict())
        
        return citations
    
    def generate_topic_visualization(self, result: TopicAnalysisResult) -> Dict[str, Any]:
        """Generar visualizaciones mejoradas para los temas"""
        try:
            # Preparar datos para visualización
            topics_data = []
            wordclouds = {}
            
            for topic in result.topics:
                # Datos básicos del tema
                topics_data.append({
                    'topic_id': topic.topic_id,
                    'topic_name': topic.topic_name,
                    'frequency': topic.frequency,
                    'coherence': topic.coherence_score,
                    'confidence': topic.confidence,
                    'keywords': ', '.join(topic.keywords[:5])
                })
                
                # Generar nube de palabras para cada tema
                try:
                    wordcloud = self._generate_wordcloud_for_topic(topic)
                    wordclouds[topic.topic_id] = wordcloud
                except Exception as e:
                    print(f"⚠️ Error generando nube de palabras para tema {topic.topic_id}: {e}")
                    wordclouds[topic.topic_id] = None
            
            # Crear gráfico de distribución de temas mejorado
            fig_distribution = px.bar(
                topics_data,
                x='topic_name',
                y='frequency',
                title='Distribución de Frecuencia por Tema',
                labels={'frequency': 'Frecuencia', 'topic_name': 'Tema'},
                color='frequency',
                color_continuous_scale='Blues'
            )
            fig_distribution.update_layout(
                xaxis_tickangle=-45,
                height=400,
                template="plotly_dark"
            )
            
            # Crear gráfico de coherencia mejorado
            fig_coherence = px.bar(
                topics_data,
                x='topic_name',
                y='coherence',
                title='Coherencia de Temas',
                labels={'coherence': 'Coherencia', 'topic_name': 'Tema'},
                color='coherence',
                color_continuous_scale='Viridis'
            )
            fig_coherence.update_layout(
                xaxis_tickangle=-45,
                height=400,
                template="plotly_dark"
            )
            
            # Crear gráfico de confianza
            fig_confidence = px.bar(
                topics_data,
                x='topic_name',
                y='confidence',
                title='Confianza de Temas',
                labels={'confidence': 'Confianza', 'topic_name': 'Tema'},
                color='confidence',
                color_continuous_scale='Greens'
            )
            fig_confidence.update_layout(
                xaxis_tickangle=-45,
                height=400,
                template="plotly_dark"
            )
            
            return {
                'distribution_chart': fig_distribution,
                'coherence_chart': fig_coherence,
                'confidence_chart': fig_confidence,
                'wordclouds': wordclouds,
                'topics_data': topics_data,
                'algorithm': result.algorithm_used,
                'total_topics': result.total_topics,
                'avg_coherence': result.coherence_avg
            }
            
        except Exception as e:
            logger.error(f"Error generando visualizaciones: {str(e)}")
            return {}
    
    def _generate_wordcloud_for_topic(self, topic: TopicData) -> Optional[str]:
        """Generar nube de palabras para un tema específico"""
        try:
            from wordcloud import WordCloud
            import matplotlib.pyplot as plt
            import matplotlib
            import io
            import base64
            
            # Configurar matplotlib para mejor compatibilidad con Streamlit
            matplotlib.use('Agg')  # Usar backend sin GUI
            
            # Preparar texto para nube de palabras
            # Usar palabras clave con pesos basados en frecuencia
            word_freq = {}
            for i, keyword in enumerate(topic.keywords[:15]):  # Top 15 palabras
                # Peso decreciente: primera palabra = peso alto, última = peso bajo
                weight = len(topic.keywords) - i
                word_freq[keyword] = weight
            
            if not word_freq:
                return None
            
            # Crear nube de palabras con configuración optimizada para Streamlit
            wordcloud = WordCloud(
                width=400,
                height=300,
                background_color='white',  # Fondo blanco para mejor visibilidad
                colormap='viridis',
                max_words=20,
                relative_scaling=0.5,
                random_state=42,
                prefer_horizontal=0.9,  # Preferir texto horizontal
                min_font_size=10,
                max_font_size=60
            ).generate_from_frequencies(word_freq)
            
            # Convertir a imagen base64
            fig, ax = plt.subplots(figsize=(6, 4))
            ax.imshow(wordcloud, interpolation='bilinear')
            ax.axis('off')
            ax.set_title(f'{topic.topic_name}', 
                        fontsize=12, color='black', pad=10)
            
            # Guardar en buffer con configuración optimizada
            buffer = io.BytesIO()
            fig.savefig(buffer, format='png', bbox_inches='tight', 
                       facecolor='white', edgecolor='none', dpi=100)
            buffer.seek(0)
            
            # Convertir a base64
            img_base64 = base64.b64encode(buffer.getvalue()).decode()
            plt.close(fig)
            
            return f"data:image/png;base64,{img_base64}"
            
        except ImportError:
            print("⚠️ WordCloud no está disponible. Instala con: pip install wordcloud")
            return None
        except Exception as e:
            print(f"⚠️ Error generando nube de palabras: {e}")
            return None
    
    def calculate_topic_relationships(self, result: TopicAnalysisResult) -> Dict[str, Any]:
        """Calcular relaciones entre temas de manera precisa"""
        try:
            topics = result.topics
            n_topics = len(topics)
            
            if n_topics < 2:
                return {
                    'similarity_matrix': [],
                    'similar_topics': [],
                    'shared_words': {}
                }
            
            # Crear matriz de similitud
            similarity_matrix = np.zeros((n_topics, n_topics))
            shared_words = {}
            
            for i in range(n_topics):
                for j in range(n_topics):
                    if i == j:
                        similarity_matrix[i][j] = 1.0
                    else:
                        # Calcular similitud basada en palabras clave compartidas
                        topic1_keywords = set(topics[i].keywords[:10])
                        topic2_keywords = set(topics[j].keywords[:10])
                        
                        # Intersección de palabras clave
                        intersection = topic1_keywords & topic2_keywords
                        union = topic1_keywords | topic2_keywords
                        
                        if len(union) > 0:
                            # Coeficiente de Jaccard
                            jaccard_similarity = len(intersection) / len(union)
                            similarity_matrix[i][j] = jaccard_similarity
                            
                            # Guardar palabras compartidas
                            shared_words[f"{i}_{j}"] = list(intersection)
                        else:
                            similarity_matrix[i][j] = 0.0
            
            # Encontrar temas más similares
            similar_topics = []
            for i in range(n_topics):
                for j in range(i + 1, n_topics):
                    similarity = similarity_matrix[i][j]
                    if similarity > 0.1:  # Umbral mínimo de similitud
                        similar_topics.append({
                            'topic1_id': i,
                            'topic1_name': topics[i].topic_name,
                            'topic2_id': j,
                            'topic2_name': topics[j].topic_name,
                            'similarity': similarity,
                            'shared_words': shared_words.get(f"{i}_{j}", [])
                        })
            
            # Ordenar por similitud descendente
            similar_topics.sort(key=lambda x: x['similarity'], reverse=True)
            
            return {
                'similarity_matrix': similarity_matrix.tolist(),
                'similar_topics': similar_topics,
                'shared_words': shared_words,
                'topic_names': [topic.topic_name for topic in topics]
            }
            
        except Exception as e:
            print(f"⚠️ Error calculando relaciones entre temas: {e}")
            return {
                'similarity_matrix': [],
                'similar_topics': [],
                'shared_words': {}
            }
    
    def generate_word_document(self, result: TopicAnalysisResult, title: str = "Análisis de Temas", 
                              author: str = "Investigador") -> bytes:
        """Generar documento Word con los temas analizados"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
            import io
            from datetime import datetime
        except ImportError:
            raise ImportError("python-docx es requerido para generar documentos Word. Instala con: pip install python-docx")
        
        # Crear documento
        doc = Document()
        
        # Configurar estilos
        title_style = doc.styles['Title']
        title_style.font.size = Pt(18)
        title_style.font.bold = True
        
        heading_style = doc.styles['Heading 1']
        heading_style.font.size = Pt(14)
        heading_style.font.bold = True
        
        # Título del documento
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Información del documento
        doc.add_paragraph(f"Autor: {author}")
        doc.add_paragraph(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph("")
        
        # Resumen ejecutivo
        doc.add_heading("Resumen Ejecutivo", level=1)
        
        summary_text = f"""
        Este documento presenta un análisis cualitativo de temas extraídos de documentos de investigación.
        
        Estadísticas del análisis:
        • Total de temas identificados: {result.total_topics}
        • Algoritmo utilizado: {result.algorithm_used}
        • Coherencia promedio: {result.coherence_avg:.3f}
        • Tiempo de procesamiento: {result.processing_time:.2f} segundos
        """
        
        doc.add_paragraph(summary_text)
        doc.add_paragraph("")
        
        # Temas principales
        doc.add_heading("Temas Identificados", level=1)
        
        for i, topic in enumerate(result.topics, 1):
            # Título del tema
            topic_title = f"{i}. {topic.topic_name}"
            doc.add_heading(topic_title, level=2)
            
            # Información básica
            info_text = f"""
            Frecuencia: {topic.frequency} apariciones
            Coherencia: {topic.coherence_score:.3f}
            Confianza: {topic.confidence:.3f}
            Documentos relacionados: {len(topic.documents)}
            """
            
            doc.add_paragraph(info_text)
            
            # Palabras clave
            doc.add_heading("Palabras Clave", level=3)
            keywords_text = ", ".join(topic.keywords[:10])
            doc.add_paragraph(keywords_text)
            
            # Descripción del tema
            if topic.description:
                doc.add_heading("Descripción", level=3)
                doc.add_paragraph(topic.description)
            
            # Documentos relacionados
            if topic.documents:
                doc.add_heading("Documentos Relacionados", level=3)
                for doc_name in topic.documents:
                    doc.add_paragraph(f"• {doc_name}")
            
            doc.add_paragraph("")  # Espacio entre temas
        
        # Análisis de relaciones
        relationships = self.calculate_topic_relationships(result)
        if relationships['similar_topics']:
            doc.add_heading("Relaciones entre Temas", level=1)
            
            for rel in relationships['similar_topics']:
                rel_text = f"""
                {rel['topic1_name']} ↔ {rel['topic2_name']}
                Similitud: {rel['similarity']:.3f}
                Palabras compartidas: {', '.join(rel['shared_words']) if rel['shared_words'] else 'Ninguna'}
                """
                doc.add_paragraph(rel_text)
        
        # Guardar en bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        return doc_bytes.getvalue()
    
    def export_topic_analysis(self, result: TopicAnalysisResult, format: str = 'json') -> Dict[str, Any]:
        """Exportar análisis de temas"""
        try:
            export_data = {
                'metadata': {
                    'algorithm': result.algorithm_used,
                    'total_topics': result.total_topics,
                    'avg_coherence': result.coherence_avg,
                    'processing_time': result.processing_time,
                    'export_date': datetime.now().isoformat()
                },
                'topics': []
            }
            
            for topic in result.topics:
                topic_data = {
                    'topic_id': topic.topic_id,
                    'topic_name': topic.topic_name,
                    'keywords': topic.keywords,
                    'coherence_score': topic.coherence_score,
                    'frequency': topic.frequency,
                    'confidence': topic.confidence,
                    'description': topic.description,
                    'documents': topic.documents,
                    'citations_count': len(topic.citations)
                }
                export_data['topics'].append(topic_data)
            
            return export_data
            
        except Exception as e:
            logger.error(f"Error exportando análisis: {str(e)}")
            return {}
