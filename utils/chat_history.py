"""\nUtilidades para el historial de conversación y exportación\n"""

import json
import os
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Optional
import streamlit as st
from io import BytesIO

# Importaciones para exportación
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

from config.settings import config
from utils.logger import setup_logger

logger = setup_logger("ChatHistory")

class ChatHistoryManager:
    """Gestor del historial de conversación"""
    
    def __init__(self):
        self.history_dir = config.DATA_DIR / "chat_history"
        self.history_dir.mkdir(exist_ok=True)
        
    def save_conversation(self, messages: List[Dict], conversation_name: Optional[str] = None) -> str:
        """\n        Guardar conversación en formato JSON\n        \n        Args:\n            messages: Lista de mensajes de la conversación\n            conversation_name: Nombre personalizado para la conversación\n            \n        Returns:\n            str: Nombre del archivo guardado\n        """
        try:
            if not conversation_name:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                conversation_name = f"conversacion_{timestamp}"
            
            # Sanitizar el nombre del archivo para Windows
            import re
            # Reemplazar caracteres inválidos con guiones bajos
            safe_conversation_name = re.sub(r'[<>:"/\\|?*]', '_', conversation_name)
            # Eliminar espacios múltiples y reemplazar con guiones bajos
            safe_conversation_name = re.sub(r'\s+', '_', safe_conversation_name)
            # Limitar la longitud del nombre
            safe_conversation_name = safe_conversation_name[:100]
            
            # Estructura del historial
            conversation_data = {
                "metadata": {
                    "name": conversation_name,  # Nombre original
                    "safe_filename": safe_conversation_name,  # Nombre seguro para archivo
                    "created_at": datetime.now().isoformat(),
                    "total_messages": len(messages),
                    "user_messages": len([m for m in messages if m["role"] == "user"]),
                    "assistant_messages": len([m for m in messages if m["role"] == "assistant"]),
                    "version": "1.0"
                },
                "messages": messages,
                "statistics": {
                    "rag_responses": len([m for m in messages if m["role"] == "assistant" and m.get("context_used")]),
                    "models_used": list(set([m.get("model_used", "unknown") for m in messages if m["role"] == "assistant"])),
                    "total_tokens_estimated": sum([len(m["content"]) // 4 for m in messages])  # Estimación aproximada
                }
            }
            
            filename = f"{safe_conversation_name}.json"
            filepath = self.history_dir / filename
            
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(conversation_data, f, indent=2, ensure_ascii=False)
            
            logger.info(f"Conversación guardada: {filename}")
            return filename
            
        except Exception as e:
            logger.error(f"Error al guardar conversación: {e}")
            raise
    
    def load_conversation(self, filename: str) -> Optional[Dict]:
        """\n        Cargar conversación desde archivo JSON\n        \n        Args:\n            filename: Nombre del archivo a cargar\n            \n        Returns:\n            Dict: Datos de la conversación o None si hay error\n        """
        try:
            filepath = self.history_dir / filename
            
            if not filepath.exists():
                logger.warning(f"Archivo no encontrado: {filename}")
                return None
            
            with open(filepath, 'r', encoding='utf-8') as f:
                conversation_data = json.load(f)
            
            logger.info(f"Conversación cargada: {filename}")
            return conversation_data
            
        except Exception as e:
            logger.error(f"Error al cargar conversación: {e}")
            return None
    
    def list_conversations(self) -> List[Dict]:
        """\n        Listar todas las conversaciones guardadas\n        \n        Returns:\n            List[Dict]: Lista de metadatos de conversaciones\n        """
        try:
            conversations = []
            
            for filepath in self.history_dir.glob("*.json"):
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                    
                    conversations.append({
                        "filename": filepath.name,
                        "name": data.get("metadata", {}).get("name", filepath.stem),
                        "created_at": data.get("metadata", {}).get("created_at"),
                        "total_messages": data.get("metadata", {}).get("total_messages", 0),
                        "file_size": filepath.stat().st_size
                    })
                except Exception as e:
                    logger.warning(f"Error al leer {filepath.name}: {e}")
                    continue
            
            # Ordenar por fecha de creación (más reciente primero)
            conversations.sort(key=lambda x: x.get("created_at", ""), reverse=True)
            return conversations
            
        except Exception as e:
            logger.error(f"Error al listar conversaciones: {e}")
            return []
    
    def delete_conversation(self, filename: str) -> bool:
        """\n        Eliminar una conversación\n        \n        Args:\n            filename: Nombre del archivo a eliminar\n            \n        Returns:\n            bool: True si se eliminó correctamente\n        """
        try:
            filepath = self.history_dir / filename
            
            if filepath.exists():
                filepath.unlink()
                logger.info(f"Conversación eliminada: {filename}")
                return True
            else:
                logger.warning(f"Archivo no encontrado para eliminar: {filename}")
                return False
                
        except Exception as e:
            logger.error(f"Error al eliminar conversación: {e}")
            return False

class ChatExporter:
    """Exportador de mensajes de chat"""
    
    def __init__(self):
        self.temp_dir = config.DATA_DIR / "temp_exports"
        self.temp_dir.mkdir(exist_ok=True)
    
    def export_message_to_docx(self, message: Dict, include_metadata: bool = True) -> Optional[BytesIO]:
        """\n        Exportar un mensaje individual a formato DOCX\n        \n        Args:\n            message: Diccionario con los datos del mensaje\n            include_metadata: Si incluir metadatos como timestamp\n            \n        Returns:\n            BytesIO: Archivo DOCX en memoria\n        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx no está instalado. Instala con: pip install python-docx")
        
        try:
            doc = Document()
            
            # Título del documento
            title = doc.add_heading('Mensaje de CogniChat', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Información del mensaje
            if include_metadata:
                doc.add_heading('Información del Mensaje', level=1)
                
                info_table = doc.add_table(rows=0, cols=2)
                info_table.style = 'Table Grid'
                
                # Agregar filas de información
                info_data = [
                    ('Rol', 'Usuario' if message['role'] == 'user' else 'Asistente'),
                    ('Timestamp', message.get('timestamp', 'N/A')),
                    ('Modelo usado', message.get('model_used', 'N/A')),
                    ('Tokens máximos', str(message.get('max_tokens', 'N/A'))),
                    ('RAG habilitado', 'Sí' if message.get('context_used') else 'No')
                ]
                
                for key, value in info_data:
                    row = info_table.add_row()
                    row.cells[0].text = key
                    row.cells[1].text = value
                
                doc.add_paragraph()  # Espacio
            
            # Contenido del mensaje
            doc.add_heading('Contenido', level=1)
            content_para = doc.add_paragraph(message['content'])
            
            # Contexto usado (si existe)
            if message.get('context_used') and include_metadata:
                doc.add_heading('Contexto RAG Utilizado', level=1)
                context_para = doc.add_paragraph(message['context_used'][:1000] + '...' if len(message['context_used']) > 1000 else message['context_used'])
                context_para.style = 'Intense Quote'
            
            # Guardar en memoria
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return buffer
            
        except Exception as e:
            logger.error(f"Error al exportar mensaje a DOCX: {e}")
            return None
    
    def export_message_to_pdf(self, message: Dict, include_metadata: bool = True) -> Optional[BytesIO]:
        """
        Exportar un mensaje individual a formato PDF
        
        Args:
            message: Diccionario con los datos del mensaje
            include_metadata: Si incluir metadatos como timestamp
            
        Returns:
            BytesIO: Archivo PDF en memoria
        """
        if not PDF_AVAILABLE:
            raise ImportError("reportlab no está instalado. Instala con: pip install reportlab")
        
        try:
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            
            # Estilos
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=1,  # Centrado
                textColor=HexColor('#2c3e50')
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                spaceBefore=20,
                spaceAfter=10,
                textColor=HexColor('#34495e')
            )
            
            content_style = ParagraphStyle(
                'CustomContent',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=12,
                leftIndent=20
            )
            
            # Función para limpiar texto y evitar errores con caracteres especiales
            def clean_text(text):
                if not text:
                    return ""
                text = str(text)
                # Escapar caracteres especiales para XML/HTML
                text = text.replace('&', '&amp;')
                text = text.replace('<', '&lt;')
                text = text.replace('>', '&gt;')
                text = text.replace('"', '&quot;')
                text = text.replace("'", '&#39;')
                # Remover caracteres de control que pueden causar problemas
                import re
                text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
                return text
            
            # Contenido del PDF
            story = []
            
            # Título
            story.append(Paragraph("Mensaje de CogniChat", title_style))
            story.append(Spacer(1, 20))
            
            # Información del mensaje
            if include_metadata:
                story.append(Paragraph("Información del Mensaje", heading_style))
                
                role_text = 'Usuario' if message.get('role') == 'user' else 'Asistente'
                timestamp = clean_text(message.get('timestamp', 'N/A'))
                model_used = clean_text(message.get('model_used', 'N/A'))
                max_tokens = clean_text(str(message.get('max_tokens', 'N/A')))
                rag_enabled = 'Sí' if message.get('context_used') else 'No'
                
                info_text = f"""
                <b>Rol:</b> {role_text}<br/>
                <b>Timestamp:</b> {timestamp}<br/>
                <b>Modelo usado:</b> {model_used}<br/>
                <b>Tokens máximos:</b> {max_tokens}<br/>
                <b>RAG habilitado:</b> {rag_enabled}
                """
                
                story.append(Paragraph(info_text, content_style))
                story.append(Spacer(1, 20))
            
            # Contenido del mensaje
            story.append(Paragraph("Contenido", heading_style))
            
            # Limpiar y dividir contenido en párrafos
            content = clean_text(message.get('content', ''))
            content_paragraphs = content.split('\n')
            
            for para in content_paragraphs:
                para = para.strip()
                if para:
                    # Limitar longitud de párrafos muy largos
                    if len(para) > 2000:
                        para = para[:2000] + '...'
                    try:
                        story.append(Paragraph(para, content_style))
                    except Exception as para_error:
                        logger.warning(f"Error con párrafo, usando texto plano: {para_error}")
                        story.append(Paragraph("[Contenido con caracteres especiales no soportados]", content_style))
            
            story.append(Spacer(1, 20))
            
            # Contexto usado (si existe)
            if message.get('context_used') and include_metadata:
                story.append(Paragraph("Contexto RAG Utilizado", heading_style))
                context_text = clean_text(str(message['context_used']))
                if len(context_text) > 1000:
                    context_text = context_text[:1000] + '...'
                try:
                    story.append(Paragraph(context_text, content_style))
                except Exception:
                    story.append(Paragraph("[Contexto con caracteres especiales no soportados]", content_style))
            
            # Construir PDF
            doc.build(story)
            buffer.seek(0)
            
            logger.info("PDF generado exitosamente")
            return buffer
            
        except Exception as e:
            logger.error(f"Error al exportar mensaje a PDF: {e}")
            import traceback
            logger.error(f"Traceback completo: {traceback.format_exc()}")
            return None
    
    def export_conversation_to_docx(self, messages: List[Dict]) -> Optional[BytesIO]:
        """\n        Exportar conversación completa a formato DOCX\n        \n        Args:\n            messages: Lista de mensajes de la conversación\n            \n        Returns:\n            BytesIO: Archivo DOCX en memoria\n        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx no está instalado. Instala con: pip install python-docx")
        
        try:
            doc = Document()
            
            # Título del documento
            title = doc.add_heading('Conversación de CogniChat', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Información general
            doc.add_heading('Información General', level=1)
            
            info_table = doc.add_table(rows=0, cols=2)
            info_table.style = 'Table Grid'
            
            info_data = [
                ('Fecha de exportación', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
                ('Total de mensajes', str(len(messages))),
                ('Mensajes de usuario', str(len([m for m in messages if m['role'] == 'user']))),
                ('Respuestas del asistente', str(len([m for m in messages if m['role'] == 'assistant']))),
                ('Respuestas con RAG', str(len([m for m in messages if m['role'] == 'assistant' and m.get('context_used')])))
            ]
            
            for key, value in info_data:
                row = info_table.add_row()
                row.cells[0].text = key
                row.cells[1].text = value
            
            doc.add_page_break()
            
            # Mensajes
            for i, message in enumerate(messages, 1):
                role_text = 'Usuario' if message['role'] == 'user' else 'Asistente'
                timestamp = message.get('timestamp', 'N/A')
                
                doc.add_heading(f'Mensaje {i} - {role_text} ({timestamp})', level=2)
                
                content_para = doc.add_paragraph(message['content'])
                
                if message['role'] == 'assistant' and message.get('context_used'):
                    doc.add_paragraph('Contexto RAG utilizado:', style='Intense Quote')
                    context_para = doc.add_paragraph(message['context_used'][:500] + '...' if len(message['context_used']) > 500 else message['context_used'])
                    context_para.style = 'Quote'
                
                doc.add_paragraph()  # Espacio entre mensajes
            
            # Guardar en memoria
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return buffer
            
        except Exception as e:
            logger.error(f"Error al exportar conversación a DOCX: {e}")
            return None
    
    def get_message_text_for_clipboard(self, message: Dict, include_metadata: bool = True) -> str:
        """\n        Obtener texto formateado de un mensaje para copiar al portapapeles\n        \n        Args:\n            message: Diccionario con los datos del mensaje\n            include_metadata: Si incluir metadatos\n            \n        Returns:\n            str: Texto formateado\n        """
        try:
            lines = []
            
            if include_metadata:
                lines.append("=" * 50)
                lines.append("MENSAJE DE COGNICHAT")
                lines.append("=" * 50)
                lines.append(f"Rol: {'Usuario' if message['role'] == 'user' else 'Asistente'}")
                lines.append(f"Timestamp: {message.get('timestamp', 'N/A')}")
                
                if message['role'] == 'assistant':
                    lines.append(f"Modelo usado: {message.get('model_used', 'N/A')}")
                    lines.append(f"RAG habilitado: {'Sí' if message.get('context_used') else 'No'}")
                
                lines.append("-" * 50)
                lines.append("CONTENIDO:")
                lines.append("-" * 50)
            
            lines.append(message['content'])
            
            if message.get('context_used') and include_metadata:
                lines.append("")
                lines.append("-" * 50)
                lines.append("CONTEXTO RAG UTILIZADO:")
                lines.append("-" * 50)
                context_text = message['context_used'][:500] + '...' if len(message['context_used']) > 500 else message['context_used']
                lines.append(context_text)
            
            if include_metadata:
                lines.append("")
                lines.append("=" * 50)
            
            return "\n".join(lines)
            
        except Exception as e:
            logger.error(f"Error al formatear mensaje para portapapeles: {e}")
            return message.get('content', '')

# Instancias globales
chat_history_manager = ChatHistoryManager()
chat_exporter = ChatExporter()